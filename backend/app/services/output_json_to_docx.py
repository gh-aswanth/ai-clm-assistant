"""Build ``contract_review.docx`` from ``output.json`` + ``cpwd.txt``.

**Pipeline overview**

1. Read the base contract text as a list of lines (`cpwd_lines`).
2. For each **finding** in ``output.json``, locate ``document_content`` inside that list (or treat it
   as a **placeholder** for pure inserts — see ``_resolve_block_start``).
3. Emit any **unprocessed** source lines as normal paragraphs (no track changes).
4. Emit this finding’s redlines: tracked ``w:del``/``w:ins`` via ``docx_track_changes``.
5. Optionally attach a **review comment** spanning the edited range (balloon-friendly placement).
6. After all findings, append any **remaining** tail of ``cpwd.txt``.

**Finding shapes**

- ``action`` ≈ ``replaced`` (default): ``document_content`` must appear in ``cpwd.txt`` once, in order.
  Algorithm: ``difflib.SequenceMatcher`` lines from ``document_content`` → ``redline_recommendation_text``.
  Deletions get red run color ``COLOR_REPLACED``; insertions use default (black after Accept).
- ``action`` ∈ {``added``, ``add``, …}: same anchor rules, but **new** text is emitted as empty
  delete + insert per line. If ``document_content`` is a **placeholder** not in ``cpwd`` (e.g.
  ``NEW CLAUSE TO BE INSERTED…``), only ``redline_recommendation_text`` is written (no placeholder
  paragraphs).

**Requires:** ``pip install python-docx``
"""

from __future__ import annotations
import io

import argparse
import difflib
import logging
import json
import warnings
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from .docx_track_changes import (
    REVISION_ID_BASE,
    append_tracked_replacement_paragraph,
    ensure_track_revisions,
    runs_in_paragraph_doc_order,
    utc_revision_timestamp,
)

ROOT = Path(__file__).resolve().parent
DEFAULT_JSON = ROOT / "output.json"
DEFAULT_SOURCE = ROOT / "cpwd.txt"
DEFAULT_OUT = ROOT / "contract_review.docx"

logger = logging.getLogger(__name__)
AUTHOR = "AI Assistant"
INITIALS = "Assistant"
# Word ``w:color`` RR BB GG for struck deletions (optional; None would use Word default).
COLOR_REPLACED = "C00000"


def _line_empty(line: str) -> bool:
    """True if ``line`` is None, whitespace-only, or empty (after strip).

    **Example:** ``\"  \\n\"`` → True; ``\"\\u200b\"`` (zero-width space) → False — used to keep a
    tracked “blank” paragraph visible to Word.
    """
    return not (line or "").strip()


def _cpwd_lines(text: str) -> list[str]:
    """Split source file into lines and trim **leading** spaces only per line.

    **Why leading-only:** Contracts may use indentation in ``cpwd.txt``; we preserve trailing spacing
    and **do not** drop trailing blank lines (line indices must stay aligned with ``_find_block``).

    **Example:** ``\"    Hello\\n\"`` → ``[\"Hello\"]``.
    """
    return [ln.lstrip() for ln in text.splitlines()]


def _trimmed_paragraph_lines(text: str) -> list[str]:
    """Normalize a JSON **multi-line string** field into logical paragraphs.

    **Steps:** Strip outer whitespace; split on ``\\n``; lstrip each line; drop trailing empty lines.

    **Example:** ``\"\\n  Alpha\\n\\n\"`` → ``[\"Alpha\"]``. Used for ``document_content`` and redline text.
    """
    lines = [ln.lstrip() for ln in text.strip().split("\n")]
    while lines and lines[-1] == "":
        lines.pop()
    return lines


def _normalize_for_match(line: str) -> str:
    """Collapse internal whitespace so ``cpwd`` vs JSON comparison tolerates spacing.

    **Example:** ``\"A   B\"`` and ``\"A B\"`` compare equal.
    """
    return " ".join(line.split())


def _find_block(cpwd: list[str], needle: list[str], start: int) -> int:
    """Find the first index ``i >= start`` where ``needle`` lines match ``cpwd[i : i+len(needle)]``.

    **Matching:** Line-by-line after ``_normalize_for_match`` (not raw string equality).

    **Args:** ``needle`` empty → returns ``start`` (insert at cursor with no source block).

    **Raises:** ``ValueError`` if no window matches (with a short preview of ``needle[0]``).

    **Example:** ``cpwd = [\"A\",\"B\",\"C\"]``, ``needle = [\"B\"]``, ``start=1`` → ``1``.
    """
    if not needle:
        return start
    n = len(needle)
    lead = _normalize_for_match(needle[0])
    for i in range(start, len(cpwd) - n + 1):
        if _normalize_for_match(cpwd[i]) != lead:
            continue
        if all(
            _normalize_for_match(cpwd[i + j]) == _normalize_for_match(needle[j])
            for j in range(n)
        ):
            return i
    preview = needle[0][:80]
    raise ValueError(
        f"Could not locate finding block starting near {preview!r} (search from line index {start})."
    )


def _resolve_block_start(
    cpwd: list[str], needle: list[str], cursor: int, *, added_mode: bool
) -> tuple[int | None, bool]:
    """Resolve where the current finding starts in ``cpwd`` and whether ``needle`` is a placeholder.

    **Returns**

    - ``(index, False)`` — Normal case: block found at ``index`` (either from ``cursor`` forward or
      from 0 when the forward search failed but a later duplicate exists).
    - ``(cursor, True)`` — **Placeholder insert:** ``needle`` never appears in ``cpwd`` *and*
      ``added_mode`` is true (e.g. ``document_content``: ``"NEW CLAUSE TO INSERT…"``). Emit only the
      redline text; do not treat ``needle`` as real source lines.
    - ``(None, False)`` — Block exists **only** before ``cursor`` (already emitted); caller should
      skip or warn (duplicate / wrong JSON order).

    **Example (placeholder):** ``needle = [\"NO SUCH LINE\"]``, ``added_mode=True``, ``cursor=10`` →
    ``(10, True)`` so new text is inserted at the current read position without searching ``cpwd``.

    **Example (duplicate):** Same real block twice, second time ``find`` from ``cursor`` fails and
    global match has ``j < cursor`` → ``(None, False)``.
    """
    if not needle:
        return cursor, False
    try:
        return _find_block(cpwd, needle, cursor), False
    except ValueError:
        pass
    try:
        j = _find_block(cpwd, needle, 0)
    except ValueError:
        if added_mode:
            return cursor, True
        return None, False
    return j, False


def _emit_plain_lines(doc: Document, lines: list[str], a: int, b: int) -> None:
    """Append paragraphs for ``lines[a:]`` up to but not including index ``b``.

    **Effect:** Empty string → empty ``w:p``; else one ``w:p`` with a single plain ``w:r``.

    **Example:** Emit contract **between** last finding end and next finding start (no track change).
    """
    for k in range(a, b):
        if lines[k] == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph().add_run(lines[k])


def _is_added_action(action: object) -> bool:
    """Return True if JSON ``action`` means “added / insert” semantics for this pipeline."""
    if not isinstance(action, str):
        return False
    return action.strip().lower() in {"add", "added", "insert", "inserted", "addition"}


def _touch_anchor(p_elm, anchors: list) -> None:
    """Update ``anchors`` with first/last ``(w:r, w:p)`` for a Word comment range.

    **Mutates:** ``anchors`` is ``[first_tuple, last_tuple]`` or ``[None, None]`` initially.

    **Skips:** Paragraphs with no runs (shouldn’t happen after a successful text emit).
    """
    runs = runs_in_paragraph_doc_order(p_elm)
    if not runs:
        return
    if anchors[0] is None:
        anchors[0] = (runs[0], p_elm)
    anchors[1] = (runs[-1], p_elm)


def _emit_tracked_line(
    doc: Document,
    orig: str,
    new: str,
    *,
    author: str,
    ts: str,
    next_rev: list[int],
    anchors: list,
    del_color: str | None,
) -> None:
    """Emit one logical line of change: blank, unchanged plain, or del+ins pair.

    **Branches**

    1. Both sides empty → blank paragraph (spacing / alignment).
    2. ``orig == new`` → plain paragraph (no revision XML).
    3. Else → ``append_tracked_replacement_paragraph`` using ``next_rev`` (consumed two ids).

    **Insertion coloring:** ``ins_color_hex`` is always ``None`` here so Accepted text stays default
    black (no trapped green ``w:color``).

    **Example:** ``orig=\"\", new=\"New clause\"`` → empty ``w:del`` + ``w:ins`` with ``New clause``.
    """
    if _line_empty(orig) and _line_empty(new):
        doc.add_paragraph()
        return
    if orig == new:
        p = doc.add_paragraph()
        p.add_run(orig)
        _touch_anchor(p._element, anchors)
        return
    p_elm = append_tracked_replacement_paragraph(
        doc,
        orig,
        new,
        author=author,
        timestamp=ts,
        del_id=next_rev[0],
        ins_id=next_rev[0] + 1,
        del_color_hex=del_color,
        ins_color_hex=None,
    )
    next_rev[0] += 2
    _touch_anchor(p_elm, anchors)


def _emit_tracked_finding(
    doc: Document,
    orig_lines: list[str],
    suggested_text: str,
    *,
    author: str,
    ts: str,
    next_rev: list[int],
    added_mode: bool,
) -> tuple[tuple | None, tuple | None]:
    """Emit track changes for one finding’s **original** vs **suggested** line lists.

    **Added mode**

    1. Optional context: each ``orig_lines`` entry becomes a plain paragraph when
       ``document_content`` is real text (not a placeholder-only insert).
    2. Each suggested line: ``emit_pair("", line)`` → tracked pure insertion (or U+200B for an
       “empty” line slot).

    **Replaced mode**

    1. ``SequenceMatcher`` opcodes drive ``emit_pair``: equal → plain when identical; delete →
       struck original + zwsp; insert → empty del + new line; replace → align heads then delete/insert
       tails.

    **Returns:** ``(first_anchor, last_anchor)`` for comment span, or ``(None, None)`` if no runs.

    **Example (replaced):** ``[\"A\",\"B\"]`` → ``[\"A\",\"X\"]`` → ``A`` unchanged, ``B`` deleted +
    ``X`` inserted.
    """
    sug_lines = _trimmed_paragraph_lines(suggested_text or "")
    anchors: list = [None, None]

    def emit_pair(o: str, s: str) -> None:
        eo, es = _line_empty(o), _line_empty(s)
        if not eo and es:
            o, s, dc = o, "\u200b", COLOR_REPLACED
        elif not eo and not es and o != s:
            dc = COLOR_REPLACED
        else:
            dc = None
        _emit_tracked_line(
            doc, o, s, author=author, ts=ts, next_rev=next_rev, anchors=anchors, del_color=dc,
        )

    if added_mode:
        for o in orig_lines:
            if _line_empty(o):
                doc.add_paragraph()
                continue
            p = doc.add_paragraph()
            p.add_run(o)
            _touch_anchor(p._element, anchors)
        for s in sug_lines:
            emit_pair("", s if not _line_empty(s) else "\u200b")
        return anchors[0], anchors[1]

    matcher = difflib.SequenceMatcher(None, orig_lines, sug_lines)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for k in range(i1, i2):
                emit_pair(orig_lines[k], sug_lines[j1 + (k - i1)])
        elif tag == "delete":
            for k in range(i1, i2):
                emit_pair(orig_lines[k], "")
        elif tag == "insert":
            for k in range(j1, j2):
                emit_pair("", sug_lines[k])
        else:
            o_seg, s_seg = orig_lines[i1:i2], sug_lines[j1:j2]
            m = min(len(o_seg), len(s_seg))
            for idx in range(m):
                emit_pair(o_seg[idx], s_seg[idx])
            for idx in range(m, len(o_seg)):
                emit_pair(o_seg[idx], "")
            for idx in range(m, len(s_seg)):
                emit_pair("", s_seg[idx])

    return anchors[0], anchors[1]


def _comment_text_for_finding(finding: dict, idx: int, total: int, *, added_mode: bool) -> str:
    """Build balloon text: prefix ``[Added]`` / ``[Replaced]`` plus ``ai_comment`` or fallback."""
    base = (finding.get("ai_comment") or "").strip()
    if not base and finding.get("clauses"):
        base = ", ".join(finding["clauses"])
    if not base:
        base = f"Finding {idx + 1}/{total}"
    tag = "[Added] " if added_mode else "[Replaced] "
    return f"{tag}{base}"


def _append_paragraph_level_comment(
    doc: Document,
    first_p_elm,
    last_p_elm,
    *,
    text: str,
    author: str,
    initials: str | None,
) -> None:
    """Attach a Word comment whose range wraps from **start of first** to **end of last** paragraph.

    **Why paragraph edges:** ``w:commentRangeStart`` inside ``w:del`` breaks balloons in some Word
    versions; inserting ``commentRangeStart`` as the **first child** of ``first_p_elm`` avoids that.

    **Steps:** Create comment → insert ``w:commentRangeStart`` at top of first ``w:p`` → append
    ``w:commentRangeEnd`` + reference run after last paragraph’s content.
    """
    comment = doc.comments.add_comment(text=text, author=author, initials=initials)
    cid = comment.comment_id
    first_p_elm.insert(0, OxmlElement("w:commentRangeStart", attrs={qn("w:id"): str(cid)}))
    end_elm = OxmlElement("w:commentRangeEnd", attrs={qn("w:id"): str(cid)})
    ref_r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_pr.append(OxmlElement("w:rStyle", attrs={qn("w:val"): "CommentReference"}))
    ref_r.append(r_pr)
    ref_r.append(OxmlElement("w:commentReference", attrs={qn("w:id"): str(cid)}))
    last_p_elm.append(end_elm)
    last_p_elm.append(ref_r)


def _build_contract_document_from_findings(findings: list, cpwd_lines: list[str]) -> Document:
    """Build an in-memory ``Document`` from finding dicts and already-split ``cpwd`` lines."""
    doc = Document()
    doc.core_properties.title = "CPWD Contract — AI review (track changes)"
    ensure_track_revisions(doc)

    ts = utc_revision_timestamp()
    next_rev = [REVISION_ID_BASE]
    cursor = 0

    for idx, finding in enumerate(findings):
        action = finding.get("action") or "replaced"
        added_mode = _is_added_action(action)
        orig_block = _trimmed_paragraph_lines(finding.get("document_content") or "")
        raw_sug = finding.get("redline_recommendation_text")
        suggested = "" if raw_sug is None else str(raw_sug)

        start_idx, placeholder_insert = _resolve_block_start(
            cpwd_lines, orig_block, cursor, added_mode=added_mode
        )
        if start_idx is None:
            preview = (orig_block[0][:60] + "…") if orig_block else "(empty)"
            warnings.warn(
                f"Skipping finding {idx + 1}/{len(findings)}: {preview!r} "
                f"matches text already processed (cursor line {cursor}).",
                UserWarning,
                stacklevel=2,
            )
            continue

        _emit_plain_lines(doc, cpwd_lines, cursor, start_idx)

        context_lines = [] if placeholder_insert else orig_block
        fa, la = _emit_tracked_finding(
            doc,
            context_lines,
            suggested,
            author=AUTHOR,
            ts=ts,
            next_rev=next_rev,
            added_mode=added_mode,
        )
        if fa and la:
            _append_paragraph_level_comment(
                doc,
                fa[1],
                la[1],
                text=_comment_text_for_finding(finding, idx, len(findings), added_mode=added_mode),
                author=AUTHOR,
                initials=INITIALS,
            )

        cursor = start_idx + (0 if placeholder_insert else len(orig_block))

    _emit_plain_lines(doc, cpwd_lines, cursor, len(cpwd_lines))
    return doc


def build_contract_docx_bytes(findings: list, source_text: str) -> bytes:
    """Build a ``.docx`` from finding dicts and full contract source text; return raw file bytes."""
    cpwd_lines = _cpwd_lines(source_text)
    from pprint import pprint
    pprint(findings)
    doc = _build_contract_document_from_findings(findings, cpwd_lines)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_contract_docx_from_json(
    findings: list,
    document_text: str,
) -> io.BytesIO:

    data = document_text
    findings = findings
    cpwd_lines = _cpwd_lines(data)
    doc = _build_contract_document_from_findings(findings, cpwd_lines)
    bytes_io = io.BytesIO()
    doc.save(bytes_io)
    return bytes_io


def main() -> None:
    """CLI: ``python output_json_to_docx.py --json … --source … -o …``."""
    p = argparse.ArgumentParser(description="Build DOCX from output.json findings + cpwd.txt")
    p.add_argument("--json", type=Path, default=DEFAULT_JSON, help="Path to output.json")
    p.add_argument("--source", type=Path, default=DEFAULT_SOURCE, help="Path to cpwd.txt (base text)")
    p.add_argument("-o", "--output", type=Path, default=DEFAULT_OUT, help="Output .docx path")
    args = p.parse_args()
    out = build_contract_docx_from_json(args.json, args.source, args.output)
    logger.info("Wrote %s", out)


if __name__ == "__main__":
    main()
