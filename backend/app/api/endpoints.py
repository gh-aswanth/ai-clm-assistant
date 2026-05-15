from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, UploadFile, File, Form, Query
from fastapi.responses import FileResponse, HTMLResponse
import io
import re
import mammoth
from sqlalchemy.orm import Session, selectinload
from sqlalchemy import or_, func
from typing import List, Optional
import json
import shutil
import os
import logging
from pathlib import Path
import uuid
import base64
import tempfile
import threading
import time
from datetime import datetime
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

from app.db.database import get_db
from app.models import models
from app.schemas import schemas
from functools import lru_cache

router = APIRouter()

UPLOAD_DIR = os.getenv("UPLOAD_DIR", "data/uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)
DOCX_PREVIEW_CACHE_DIR = os.getenv("DOCX_PREVIEW_CACHE_DIR", "data/preview_cache")
os.makedirs(DOCX_PREVIEW_CACHE_DIR, exist_ok=True)
logger = logging.getLogger(__name__)

DOCUMENT_DRIVE_DIR = os.getenv("DOCUMENT_DRIVE_DIR", "data/document_drive")
DOCUMENT_DRIVE_CHUNKS_DIR = os.path.join(DOCUMENT_DRIVE_DIR, "chunks")
os.makedirs(DOCUMENT_DRIVE_DIR, exist_ok=True)
os.makedirs(DOCUMENT_DRIVE_CHUNKS_DIR, exist_ok=True)

_DRIVE_UPLOAD_SESSIONS = {}
_DRIVE_UPLOAD_LOCK = threading.Lock()
_DRIVE_UPLOAD_TTL_SECONDS = 60 * 60


def _resolve_file_meta(file_id: int) -> tuple[str, Optional[int]]:
    """Return (original_filename, folder_id) for a file by querying its own DB session."""
    from app.db.database import SessionLocal

    db = SessionLocal()
    try:
        f = db.query(models.File).filter(models.File.id == file_id).first()
        if f:
            return f.original_filename or f"file #{file_id}", f.folder_id
        return f"file #{file_id}", None
    finally:
        db.close()


def _run_chunk_file_background(file_id: int, contract_id: Optional[int] = None) -> None:
    """Load ``files`` row by id, run LLM chunking, persist ``document_chunks`` (own DB session).

    Emits ``chunking_started``, ``chunking_done``, or ``chunking_failed`` notifications
    to all connected WebSocket clients via the global :data:`notifier`.
    """
    import asyncio
    from app.db.database import SessionLocal
    from app.services.document_chunker import chunk_file
    from app.services.notifier import notifier

    filename, folder_id = _resolve_file_meta(file_id)

    notifier.broadcast_from_thread({
        "type": "chunking_started",
        "file_id": file_id,
        "contract_id": contract_id,
        "folder_id": folder_id,
        "filename": filename,
        "message": f"Chunking started for {filename}",
    })

    db = SessionLocal()
    try:
        asyncio.run(chunk_file(db, file_id))
        notifier.broadcast_from_thread({
            "type": "chunking_done",
            "file_id": file_id,
            "contract_id": contract_id,
            "folder_id": folder_id,
            "filename": filename,
            "message": f"Chunking complete for {filename}",
        })
    except Exception:
        logger.exception("Background chunking failed for file_id=%s", file_id)
        notifier.broadcast_from_thread({
            "type": "chunking_failed",
            "file_id": file_id,
            "contract_id": contract_id,
            "folder_id": folder_id,
            "filename": filename,
            "message": f"Chunking failed for {filename}",
        })
    finally:
        db.close()


def _run_compliance_background(contract_id: int, file_id: int) -> None:
    """Structured LLM compliance checks from ``document_chunks``; own DB session.

    Emits ``compliance_started``, ``compliance_done``, or ``compliance_failed`` notifications.
    """
    import asyncio
    from app.db.database import SessionLocal
    from app.services.compliance_check import extract_and_store_compliance_checks
    from app.services.notifier import notifier

    filename, folder_id = _resolve_file_meta(file_id)

    notifier.broadcast_from_thread({
        "type": "compliance_started",
        "file_id": file_id,
        "contract_id": contract_id,
        "folder_id": folder_id,
        "filename": filename,
        "message": f"Compliance check started for {filename}",
    })

    db = SessionLocal()
    try:
        asyncio.run(
            extract_and_store_compliance_checks(db, contract_id=contract_id, file_id=file_id)
        )
        notifier.broadcast_from_thread({
            "type": "compliance_done",
            "file_id": file_id,
            "contract_id": contract_id,
            "folder_id": folder_id,
            "filename": filename,
            "message": f"Compliance checks complete for {filename}",
        })
    except Exception:
        logger.exception(
            "Background compliance failed contract_id=%s file_id=%s",
            contract_id,
            file_id,
        )
        notifier.broadcast_from_thread({
            "type": "compliance_failed",
            "file_id": file_id,
            "contract_id": contract_id,
            "folder_id": folder_id,
            "filename": filename,
            "message": f"Compliance check failed for {filename}",
        })
    finally:
        db.close()


def _run_scoring_background(contract_id: int, file_id: int) -> None:
    """Run multi-agent scoring workflow in a background thread; own DB session.

    Emits ``scoring_started``, ``scoring_done``, or ``scoring_failed`` notifications.
    """
    from app.db.database import SessionLocal
    from app.services.agent_workflow import compile_graph, load_contract_guidelines
    from app.services.notifier import notifier

    filename, folder_id = _resolve_file_meta(file_id)

    notifier.broadcast_from_thread({
        "type": "scoring_started",
        "file_id": file_id,
        "contract_id": contract_id,
        "folder_id": folder_id,
        "filename": filename,
        "message": f"Scoring started for {filename}",
    })

    db = SessionLocal()
    try:
        version = (
            db.query(models.DocumentVersion)
            .filter(
                models.DocumentVersion.contract_id == contract_id,
                models.DocumentVersion.file_id == file_id,
            )
            .order_by(models.DocumentVersion.version_number.desc())
            .first()
        )
        if not version:
            raise ValueError(f"No document version found for contract_id={contract_id} file_id={file_id}")

        initial_state = load_contract_guidelines(contract_id, db)

        if not initial_state.get("contract_text"):
            chunks = (
                db.query(models.DocumentChunk)
                .filter(models.DocumentChunk.file_id == file_id)
                .order_by(models.DocumentChunk.chunk_index)
                .all()
            )
            if chunks:
                initial_state["contract_text"] = "\n\n".join(ch.content for ch in chunks)

        if not initial_state.get("contract_text"):
            raise ValueError("No contract text available for scoring")

        graph = compile_graph()
        result_state = graph.invoke(initial_state)

        result_json = {
            "contract_text": result_state.get("contract_text", ""),
            "contract_metadata": result_state.get("contract_metadata", {}),
            "classifications": result_state.get("classifications", {}),
            "validation_results": result_state.get("validation_results", []),
            "finance_analysis": result_state.get("finance_analysis", {}),
            "cross_validation_result": result_state.get("cross_validation_result", {}),
        }

        row = models.ScoringResult(
            contract_id=contract_id,
            document_version_id=version.id,
            file_id=file_id,
            result_json=result_json,
        )
        db.add(row)
        _auto_milestone(
            db, contract_id,
            title="Scoring completed",
            description=f"Automated multi-agent scoring completed for document version v{version.version_number}. File #{file_id}.",
        )
        db.commit()

        notifier.broadcast_from_thread({
            "type": "scoring_done",
            "file_id": file_id,
            "contract_id": contract_id,
            "folder_id": folder_id,
            "filename": filename,
            "message": f"Scoring complete for {filename}",
        })
    except Exception:
        logger.exception("Background scoring failed for contract_id=%s file_id=%s", contract_id, file_id)
        notifier.broadcast_from_thread({
            "type": "scoring_failed",
            "file_id": file_id,
            "contract_id": contract_id,
            "folder_id": folder_id,
            "filename": filename,
            "message": f"Scoring failed for {filename}",
        })
    finally:
        db.close()


def _run_graph_build_background(contract_id: int, file_id: int) -> None:
    """Run knowledge-graph build from document chunks in a background thread; own DB session.

    Emits ``graph_build_started``, ``graph_build_done``, or ``graph_build_failed`` notifications.
    """
    from app.db.database import SessionLocal
    from app.services.graph_builder import build_graph
    from app.services.notifier import notifier

    filename, folder_id = _resolve_file_meta(file_id)

    notifier.broadcast_from_thread({
        "type": "graph_build_started",
        "file_id": file_id,
        "contract_id": contract_id,
        "folder_id": folder_id,
        "filename": filename,
        "message": f"Graph build started for {filename}",
    })

    db = SessionLocal()
    try:
        version = (
            db.query(models.DocumentVersion)
            .filter(
                models.DocumentVersion.contract_id == contract_id,
                models.DocumentVersion.file_id == file_id,
            )
            .order_by(models.DocumentVersion.version_number.desc())
            .first()
        )
        if not version:
            raise ValueError(f"No document version found for contract_id={contract_id} file_id={file_id}")

        build_graph(db, contract_id, version.id)
        _auto_milestone(
            db, contract_id,
            title="Knowledge graph built",
            description=f"Knowledge graph built automatically for document version v{version.version_number}. File #{file_id}.",
        )
        db.commit()

        notifier.broadcast_from_thread({
            "type": "graph_build_done",
            "file_id": file_id,
            "contract_id": contract_id,
            "folder_id": folder_id,
            "filename": filename,
            "message": f"Graph build complete for {filename}",
        })
    except Exception:
        logger.exception("Background graph build failed for contract_id=%s file_id=%s", contract_id, file_id)
        notifier.broadcast_from_thread({
            "type": "graph_build_failed",
            "file_id": file_id,
            "contract_id": contract_id,
            "folder_id": folder_id,
            "filename": filename,
            "message": f"Graph build failed for {filename}",
        })
    finally:
        db.close()


def _run_post_chunk_tasks_background(contract_id: int, file_id: int) -> None:
    """Run compliance, scoring, and graph build in parallel (chunks must already exist)."""
    threads = [
        threading.Thread(target=_run_compliance_background, args=(contract_id, file_id), daemon=True),
        threading.Thread(target=_run_scoring_background, args=(contract_id, file_id), daemon=True),
        threading.Thread(target=_run_graph_build_background, args=(contract_id, file_id), daemon=True),
    ]
    for t in threads:
        t.start()
    for t in threads:
        t.join()


def _run_chunk_and_compliance_background(contract_id: int, file_id: int) -> None:
    """Chunk first, then run compliance, scoring, and graph build in parallel."""
    _run_chunk_file_background(file_id, contract_id=contract_id)
    _run_post_chunk_tasks_background(contract_id, file_id)


def _schedule_chunk_file(file_id: int, contract_id: Optional[int] = None) -> None:
    """Run :func:`_run_chunk_file_background` in a daemon thread (fallback when ``BackgroundTasks`` is unavailable)."""
    thread = threading.Thread(
        target=_run_chunk_file_background, args=(file_id,), kwargs={"contract_id": contract_id}, daemon=True
    )
    thread.start()


def _create_contract_upload_file(
    db: Session,
    *,
    disk_path: str,
    original_filename: str,
    file_type: str,
    content_type: Optional[str],
) -> models.File:
    """Persist a ``files`` row for an upload saved at ``disk_path`` (contract uploads: no ``folder_id``)."""
    size_bytes = os.path.getsize(disk_path) if os.path.isfile(disk_path) else 0
    rec = models.File(
        original_filename=original_filename,
        file_path=disk_path,
        file_type=file_type,
        content_type=content_type,
        size_bytes=size_bytes,
        folder_id=None,
    )
    db.add(rec)
    db.flush()
    return rec


def _safe_filename(name: str) -> str:
    safe = "".join(ch if ch.isalnum() or ch in ("-", "_", ".", " ") else "_" for ch in name)
    return safe.strip().replace(" ", "_") or f"file-{uuid.uuid4().hex}"


def _parse_int(value: Optional[str], name: str) -> int:
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        raise HTTPException(status_code=400, detail=f"Invalid {name}: must be an integer")
    if parsed < 0:
        raise HTTPException(status_code=400, detail=f"Invalid {name}: must be non-negative")
    return parsed


def _resolve_drive_file_path(folder_id: int, original_filename: str, upload_id: str) -> str:
    suffix = os.path.splitext(original_filename)[1]
    safe_name = _safe_filename(os.path.splitext(original_filename)[0])
    return os.path.join(DOCUMENT_DRIVE_DIR, f"{folder_id}_{upload_id}_{safe_name}{suffix}")


def _cleanup_expired_upload_sessions():
    now = time.time()
    stale = []
    with _DRIVE_UPLOAD_LOCK:
        for upload_id, session in list(_DRIVE_UPLOAD_SESSIONS.items()):
            if now - session["started_at"] > _DRIVE_UPLOAD_TTL_SECONDS:
                stale.append(upload_id)
        for upload_id in stale:
            session = _DRIVE_UPLOAD_SESSIONS.pop(upload_id, None)
            if not session:
                continue
            chunk_dir = session["chunk_dir"]
            if os.path.isdir(chunk_dir):
                try:
                    shutil.rmtree(chunk_dir, ignore_errors=True)
                except OSError:
                    pass


# ── Helpers ───────────────────────────────────────────────────────────────────

def _parse_date(date_str: str):
    if not date_str:
        return None
    try:
        return datetime.fromisoformat(date_str.replace("Z", "+00:00"))
    except ValueError:
        return datetime.strptime(date_str, "%Y-%m-%d")

def _next_contract_number(db: Session) -> str:
    """Next CLM-{year}-{seq} using max existing seq for that year (not row count)."""
    year = datetime.utcnow().year
    prefix = f"CLM-{year}-"
    pattern = re.compile(re.escape(prefix) + r"(\d+)$")
    max_seq = 0
    for (cn,) in db.query(models.Contract.contract_number).filter(
        models.Contract.contract_number.isnot(None),
        models.Contract.contract_number.like(f"{prefix}%"),
    ):
        m = pattern.match(cn)
        if m:
            max_seq = max(max_seq, int(m.group(1)))
    return f"{prefix}{max_seq + 1:04d}"

def _mark_version_latest(db: Session, contract_id: int, new_version_id: int):
    """Unset is_latest on all versions for the contract, then set on the new one."""
    db.query(models.DocumentVersion).filter(
        models.DocumentVersion.contract_id == contract_id
    ).update({"is_latest": False})
    db.query(models.DocumentVersion).filter(
        models.DocumentVersion.id == new_version_id
    ).update({"is_latest": True})


def _auto_milestone(db: Session, contract_id: int, title: str, description: str, status: str = "completed"):
    """Insert an automatic lifecycle milestone for a contract."""
    db.add(models.Milestone(
        contract_id=contract_id,
        title=title,
        description=description,
        due_date=datetime.utcnow(),
        status=status,
    ))


def _safe_unlink(path: Optional[str]) -> None:
    if path and os.path.isfile(path):
        try:
            os.remove(path)
        except OSError:
            pass


def _docx_preview_cache_path(path: str) -> str:
    stat = os.stat(path)
    base = os.path.basename(path)
    base_no_ext = os.path.splitext(base)[0]
    safe_base = "".join(ch if ch.isalnum() or ch in ("-", "_", ".") else "_" for ch in base_no_ext)
    fingerprint = f"{int(stat.st_mtime_ns)}_{stat.st_size}"
    return os.path.join(DOCX_PREVIEW_CACHE_DIR, f"{safe_base}_{fingerprint}.html")


def _read_docx_preview_cache(cache_path: str) -> Optional[str]:
    try:
        with open(cache_path, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        return None
    except OSError:
        return None


@lru_cache(maxsize=64)
def _build_cached_docx_html(path: str, file_mtime_ns: int, file_size: int) -> str:
    """Convert DOCX to HTML and cache the resulting payload in memory by file fingerprint."""
    with open(path, "rb") as f:
        result = mammoth.convert_to_html(f)
    html_body = result.value
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <style>
    * {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{
      font-family: 'Segoe UI', Arial, sans-serif;
      font-size: 13px;
      line-height: 1.75;
      color: #1a1a2e;
      background: #fff;
      padding: 48px 56px;
      max-width: 860px;
      margin: 0 auto;
    }}
    h1 {{ font-size: 1.6rem; font-weight: 700; margin: 1.2rem 0 .6rem; }}
    h2 {{ font-size: 1.3rem; font-weight: 700; margin: 1rem 0 .5rem; }}
    h3 {{ font-size: 1.1rem; font-weight: 600; margin: .8rem 0 .4rem; }}
    h4,h5,h6 {{ font-weight: 600; margin: .6rem 0 .3rem; }}
    p  {{ margin: .4rem 0; }}
    ul,ol {{ padding-left: 1.5rem; margin: .4rem 0; }}
    li {{ margin: .2rem 0; }}
    table {{ border-collapse: collapse; width: 100%; margin: .8rem 0; }}
    td,th {{ border: 1px solid #d1d5db; padding: .35rem .6rem; text-align: left; }}
    th {{ background: #f3f4f6; font-weight: 600; }}
    strong {{ font-weight: 700; }}
    em {{ font-style: italic; }}
    u  {{ text-decoration: underline; }}
    a  {{ color: #2563eb; }}
  </style>
</head>
<body>{html_body}</body>
</html>"""


# ── Templates ─────────────────────────────────────────────────────────────────

@router.post("/templates/", response_model=schemas.TemplateSchema)
def create_template(template: schemas.TemplateCreate, db: Session = Depends(get_db)):
    db_template = models.Template(**template.model_dump())
    db.add(db_template)
    db.commit()
    db.refresh(db_template)
    return db_template

@router.get("/templates/", response_model=List[schemas.TemplateSchema])
def read_templates(skip: int = 0, limit: int = 100, db: Session = Depends(get_db)):
    return db.query(models.Template).offset(skip).limit(limit).all()


# ── Master Signers ────────────────────────────────────────────────────────────

@router.get("/master-signers/", response_model=List[schemas.MasterSignerSchema])
def list_master_signers(search: Optional[str] = None, db: Session = Depends(get_db)):
    query = db.query(models.MasterSigner).filter(models.MasterSigner.is_active == True)
    if search:
        query = query.filter(
            models.MasterSigner.name.ilike(f"%{search}%") |
            models.MasterSigner.email.ilike(f"%{search}%") |
            models.MasterSigner.organization.ilike(f"%{search}%")
        )
    return query.order_by(models.MasterSigner.name).all()

@router.post("/master-signers/", response_model=schemas.MasterSignerSchema)
def create_master_signer(signer: schemas.MasterSignerCreate, db: Session = Depends(get_db)):
    existing = db.query(models.MasterSigner).filter(
        models.MasterSigner.email == signer.email
    ).first()
    if existing:
        raise HTTPException(status_code=409, detail="A signer with this email already exists")
    db_signer = models.MasterSigner(**signer.model_dump())
    db.add(db_signer)
    db.commit()
    db.refresh(db_signer)
    return db_signer

@router.patch("/master-signers/{signer_id}", response_model=schemas.MasterSignerSchema)
def update_master_signer(signer_id: int, signer: schemas.MasterSignerCreate, db: Session = Depends(get_db)):
    db_signer = db.query(models.MasterSigner).filter(models.MasterSigner.id == signer_id).first()
    if not db_signer:
        raise HTTPException(status_code=404, detail="Signer not found")
    for key, value in signer.model_dump(exclude_unset=True).items():
        setattr(db_signer, key, value)
    db.commit()
    db.refresh(db_signer)
    return db_signer

@router.delete("/master-signers/{signer_id}")
def deactivate_master_signer(signer_id: int, db: Session = Depends(get_db)):
    db_signer = db.query(models.MasterSigner).filter(models.MasterSigner.id == signer_id).first()
    if not db_signer:
        raise HTTPException(status_code=404, detail="Signer not found")
    db_signer.is_active = False
    db.commit()
    return {"detail": "Signer deactivated"}


# ── Contracts ─────────────────────────────────────────────────────────────────

@router.post("/contracts/", response_model=schemas.ContractSchema)
def create_contract(
    background_tasks: BackgroundTasks,
    title: str = Form(...),
    description: Optional[str] = Form(None),
    value: float = Form(0.0),
    start_date: str = Form(...),
    end_date: str = Form(...),
    file: Optional[UploadFile] = File(None),
    drive_file_id: Optional[int] = Form(None),
    file_id: Optional[int] = Form(None),
    guideline_framework_slug: Optional[str] = Form(None),
    guideline_rule_keys: Optional[str] = Form(None),
    db: Session = Depends(get_db),
):
    """Create a contract. When a document is provided, a ``files`` row is created first, then
    v1 ``DocumentVersion`` with ``file_id``; LLM chunking runs in a background task and writes
    ``document_chunks`` keyed by ``file_id``.
    """
    drive_file_id = file_id or drive_file_id
    db_file: Optional[models.File] = None
    uploaded_path: Optional[str] = None

    if file:
        if not file.filename:
            raise HTTPException(status_code=400, detail="Uploaded file has no filename")
        file_type = file.filename.rsplit(".", 1)[-1].lower()
        uploaded_path = os.path.join(UPLOAD_DIR, f"{datetime.now().timestamp()}_{file.filename}")
        try:
            with open(uploaded_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            db_file = _create_contract_upload_file(
                db,
                disk_path=uploaded_path,
                original_filename=file.filename,
                file_type=file_type,
                content_type=file.content_type,
            )
        except Exception as e:
            if uploaded_path and os.path.exists(uploaded_path):
                os.remove(uploaded_path)
            raise HTTPException(status_code=400, detail=f"Failed to save upload: {e}") from e

    elif drive_file_id is not None:
        db_file = db.query(models.File).filter(models.File.id == drive_file_id).first()
        if not db_file:
            raise HTTPException(status_code=404, detail="File not found")

    # Legacy contract fields mirror the primary file
    legacy_path = db_file.file_path if db_file else None
    legacy_type = (db_file.file_type if db_file else None) or (
        db_file.original_filename.rsplit(".", 1)[-1].lower() if db_file and db_file.original_filename else None
    )

    try:
        db_contract = models.Contract(
            contract_number=_next_contract_number(db),
            title=title,
            description=description,
            value=float(value),
            start_date=_parse_date(start_date),
            end_date=_parse_date(end_date),
            status=models.ContractStatus.DRAFT,
            file_path=legacy_path,
            file_type=legacy_type,
        )
    except Exception as e:
        if uploaded_path and os.path.exists(uploaded_path):
            os.remove(uploaded_path)
        raise HTTPException(status_code=400, detail=f"Error parsing contract data: {str(e)}") from e

    db.add(db_contract)
    db.flush()

    if db_file:
        db_version = models.DocumentVersion(
            contract_id=db_contract.id,
            version_number=1,
            label="Initial Upload",
            file_id=db_file.id,
            is_latest=True,
        )
        db.add(db_version)

    db.commit()
    db.refresh(db_contract)

    # Auto milestone: contract prepared
    if db_file:
        _auto_milestone(db, db_contract.id, "Contract Prepared",
                        f"Contract '{title}' was created and initial document uploaded.")
        db.commit()

    # Auto-save guideline data from the selected framework + section keys
    if guideline_framework_slug:
        fw = (
            db.query(models.GuidelineFramework)
            .filter(models.GuidelineFramework.slug == guideline_framework_slug)
            .first()
        )
        if fw:
            db_contract.guideline_framework_slug = fw.slug
            db_contract.guideline_framework_title = fw.title
            rule_keys = (
                [k.strip() for k in guideline_rule_keys.split(",") if k.strip()]
                if guideline_rule_keys
                else []
            )
            sections_q = db.query(models.GuidelineSection).filter(
                models.GuidelineSection.framework_id == fw.id,
            )
            if rule_keys:
                sections_q = sections_q.filter(
                    models.GuidelineSection.section_key.in_(rule_keys)
                )
            for sec in sections_q.all():
                col = f"guideline_{sec.section_key}"
                if hasattr(db_contract, col):
                    setattr(db_contract, col, sec.body)
            db.commit()
            db.refresh(db_contract)
            logger.info(
                "Auto-saved guideline '%s' (%d sections) on contract %d",
                fw.slug,
                sections_q.count() if rule_keys else 0,
                db_contract.id,
            )

    # Background chunking + LLM compliance (uses ``document_chunks`` after chunking)
    if db_file:
        existing = (
            db.query(models.DocumentChunk)
            .filter(models.DocumentChunk.file_id == db_file.id)
            .count()
        )
        if existing == 0:
            background_tasks.add_task(
                _run_chunk_and_compliance_background, db_contract.id, db_file.id
            )
            logger.info(
                "Scheduled chunking + compliance for file_id=%s (contract %s)",
                db_file.id,
                db_contract.id,
            )
        else:
            background_tasks.add_task(_run_post_chunk_tasks_background, db_contract.id, db_file.id)
            logger.info(
                "Scheduled compliance+scoring+graph (parallel) for file_id=%s (%d chunks exist, contract %s)",
                db_file.id,
                existing,
                db_contract.id,
            )

    return db_contract

@router.get("/dashboard/stats")
def get_dashboard_stats(db: Session = Depends(get_db)):
    total = db.query(func.count(models.Contract.id)).scalar()
    active = db.query(func.count(models.Contract.id)).filter(
        models.Contract.status == models.ContractStatus.ACTIVE
    ).scalar()
    pending_approval = db.query(func.count(models.Contract.id)).filter(
        models.Contract.status.in_([
            models.ContractStatus.REVIEW,
            models.ContractStatus.SIGNING,
        ])
    ).scalar()
    total_value = db.query(func.sum(models.Contract.value)).scalar() or 0.0

    recent_contracts = (
        db.query(models.Contract)
        .order_by(models.Contract.created_at.desc())
        .limit(5)
        .all()
    )

    upcoming_milestones = (
        db.query(models.Milestone)
        .filter(
            models.Milestone.status == "pending",
            models.Milestone.due_date >= datetime.utcnow(),
        )
        .order_by(models.Milestone.due_date.asc())
        .limit(5)
        .all()
    )

    def fmt_value(v):
        if v >= 1_000_000:
            return f"${v/1_000_000:.1f}M"
        if v >= 1_000:
            return f"${v/1_000:.0f}K"
        return f"${v:.0f}"

    def days_until(dt):
        delta = dt - datetime.utcnow()
        days = delta.days
        if days == 0:
            return "Due today"
        if days == 1:
            return "Due in 1 day"
        return f"Due in {days} days"

    return {
        "stats": {
            "total_contracts": total,
            "active_contracts": active,
            "pending_approval": pending_approval,
            "total_value": fmt_value(total_value),
        },
        "recent_contracts": [
            {
                "id": c.id,
                "title": c.title,
                "contract_number": c.contract_number,
                "status": c.status.value if c.status else "draft",
            }
            for c in recent_contracts
        ],
        "upcoming_milestones": [
            {
                "id": m.id,
                "title": m.title,
                "due_label": days_until(m.due_date),
                "status": m.status,
                "contract_id": m.contract_id,
            }
            for m in upcoming_milestones
        ],
    }


@router.get("/contracts/", response_model=List[schemas.ContractSchema])
def read_contracts(skip: int = 0, limit: int = 100, db: Session = Depends(get_db)):
    return (
        db.query(models.Contract)
        .options(
            selectinload(models.Contract.document_versions).selectinload(models.DocumentVersion.file),
        )
        .offset(skip)
        .limit(limit)
        .all()
    )

@router.get("/contracts/{contract_id}", response_model=schemas.ContractSchema)
def read_contract(contract_id: int, db: Session = Depends(get_db)):
    db_contract = (
        db.query(models.Contract)
        .options(
            selectinload(models.Contract.document_versions)
            .selectinload(models.DocumentVersion.file),
        )
        .filter(models.Contract.id == contract_id)
        .first()
    )
    if db_contract is None:
        raise HTTPException(status_code=404, detail="Contract not found")
    return db_contract


@router.delete("/contracts/{contract_id}", status_code=204)
def delete_contract(contract_id: int, db: Session = Depends(get_db)):
    db_contract = db.query(models.Contract).filter(models.Contract.id == contract_id).first()
    if not db_contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    version_file_ids = [v.file_id for v in (db_contract.document_versions or []) if v.file_id]

    for ver in list(db_contract.document_versions or []):
        db.query(models.SignatureField).filter(models.SignatureField.version_id == ver.id).delete()
        db.query(models.VersionSigner).filter(models.VersionSigner.version_id == ver.id).delete()
        db.delete(ver)

    db.query(models.ScoringResult).filter(models.ScoringResult.contract_id == contract_id).delete()
    db.query(models.ReviewItem).filter(models.ReviewItem.contract_id == contract_id).delete()
    db.query(models.SignatureField).filter(models.SignatureField.contract_id == contract_id).delete()
    db.query(models.Signer).filter(models.Signer.contract_id == contract_id).delete()
    db.query(models.ComplianceRecord).filter(models.ComplianceRecord.contract_id == contract_id).delete()
    db.query(models.Milestone).filter(models.Milestone.contract_id == contract_id).delete()

    db.delete(db_contract)
    db.flush()

    for fid in set(version_file_ids):
        if (
            db.query(models.DocumentVersion)
            .filter(models.DocumentVersion.file_id == fid)
            .count()
            == 0
        ):
            frow = db.query(models.File).filter(models.File.id == fid).first()
            if frow and frow.folder_id is None:
                db.delete(frow)

    db.commit()
    return


@router.patch("/contracts/{contract_id}/status", response_model=schemas.ContractSchema)
def update_contract_status(
    contract_id: int,
    body: dict,
    db: Session = Depends(get_db),
):
    db_contract = db.query(models.Contract).filter(models.Contract.id == contract_id).first()
    if db_contract is None:
        raise HTTPException(status_code=404, detail="Contract not found")
    new_status = body.get("status", "").lower()
    try:
        db_contract.status = models.ContractStatus(new_status)
    except ValueError:
        raise HTTPException(status_code=400, detail=f"Invalid status: {new_status}")
    _auto_milestone(
        db, contract_id,
        title=f"Status changed to {new_status.capitalize()}",
        description=f"Contract status was updated to '{new_status}'.",
    )
    db.commit()
    db.refresh(db_contract)
    return db_contract


@router.patch("/contracts/{contract_id}/guidelines", response_model=schemas.ContractSchema)
def update_contract_guidelines(
    contract_id: int,
    body: schemas.ContractGuidelineUpdate,
    db: Session = Depends(get_db),
):
    """Store guideline data on the contract — each section in its own column for AI automation."""
    db_contract = db.query(models.Contract).filter(models.Contract.id == contract_id).first()
    if db_contract is None:
        raise HTTPException(status_code=404, detail="Contract not found")
    section_fields = [
        "guideline_framework_slug", "guideline_framework_title", "guideline_snapshot",
        "guideline_financial_limits", "guideline_mandatory_clauses",
        "guideline_technical_standards", "guideline_compliance_requirements",
        "guideline_contractor_eligibility", "guideline_work_execution_standards",
        "guideline_measurement_payment", "guideline_contract_administration",
        "guideline_defect_liability", "guideline_documentation_requirements",
        "guideline_decision_thresholds", "guideline_validation_weights",
        "guideline_critical_issues",
    ]
    updated_sections = []
    for field in section_fields:
        val = getattr(body, field, None)
        if val is not None:
            setattr(db_contract, field, val)
            updated_sections.append(field.replace("guideline_", "").replace("_", " ").title())
    if updated_sections:
        _auto_milestone(
            db, contract_id,
            title="Guidelines updated",
            description=f"Guideline sections updated: {', '.join(updated_sections)}.",
        )
    db.commit()
    db.refresh(db_contract)
    return db_contract


@router.get("/contracts/{contract_id}/file")
def get_contract_file(
    contract_id: int,
    version: Optional[str] = None,
    version_id: Optional[int] = None,
    db: Session = Depends(get_db)
):
    db_contract = db.query(models.Contract).filter(models.Contract.id == contract_id).first()
    if not db_contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    path = None
    file_type = db_contract.file_type or "pdf"

    if version_id:
        db_version = db.query(models.DocumentVersion).filter(
            models.DocumentVersion.id == version_id
        ).first()
        if db_version:
            path = db_version.signed_file_path if version == "signed" else (db_version.file.file_path if db_version.file else None)
            file_type = (db_version.file.file_type if db_version.file else None)
    elif version == "signed":
        # Check latest version first, fall back to legacy
        latest = db.query(models.DocumentVersion).filter(
            models.DocumentVersion.contract_id == contract_id,
            models.DocumentVersion.is_latest == True
        ).first()
        path = (latest.signed_file_path if latest else None) or db_contract.signed_file_path
    elif version == "original":
        latest = db.query(models.DocumentVersion).filter(
            models.DocumentVersion.contract_id == contract_id,
            models.DocumentVersion.is_latest == True
        ).first()
        path = ((latest.file.file_path if latest.file else None) if latest else None) or db_contract.file_path
    else:
        # Best available: prefer signed, fall back to original
        latest = db.query(models.DocumentVersion).filter(
            models.DocumentVersion.contract_id == contract_id,
            models.DocumentVersion.is_latest == True
        ).first()
        if latest:
            path = latest.signed_file_path or (latest.file.file_path if latest.file else None)
            file_type = (latest.file.file_type if latest.file else None)
        else:
            path = db_contract.signed_file_path or db_contract.file_path

    if not path or not os.path.exists(path):
        raise HTTPException(status_code=404, detail="File not found")

    # Infer file type from path extension when the DB value is missing
    if not file_type and path:
        file_type = os.path.splitext(path)[1].lstrip(".").lower() or None

    if file_type == "pdf":
        media_type = "application/pdf"
    elif file_type in ("docx", "doc"):
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        media_type = "application/octet-stream"

    return FileResponse(path, media_type=media_type)


@router.get("/contracts/{contract_id}/preview-html")
def get_contract_preview_html(
    contract_id: int,
    version_id: Optional[int] = None,
    version: Optional[str] = None,
    db: Session = Depends(get_db),
):
    """Return DOCX file converted to HTML for in-browser preview."""
    db_contract = db.query(models.Contract).filter(models.Contract.id == contract_id).first()
    if not db_contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    path = None
    file_type = "pdf"

    if version_id:
        db_version = db.query(models.DocumentVersion).filter(
            models.DocumentVersion.id == version_id
        ).first()
        if db_version:
            path = db_version.signed_file_path if version == "signed" else (db_version.file.file_path if db_version.file else None)
            file_type = (db_version.file.file_type if db_version.file else None) or "pdf"
    else:
        latest = db.query(models.DocumentVersion).filter(
            models.DocumentVersion.contract_id == contract_id,
            models.DocumentVersion.is_latest == True,
        ).first()
        if latest:
            path = (latest.signed_file_path if version == "signed" else None) or (latest.file.file_path if latest.file else None)
            file_type = (latest.file.file_type if latest.file else None) or "pdf"
        else:
            path = (db_contract.signed_file_path if version == "signed" else None) or db_contract.file_path
            file_type = db_contract.file_type or "pdf"

    if not path or not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Signed document not yet available" if version == "signed" else "File not found")

    # Detect file type from actual path extension for signed copies (may differ from version.file_type)
    actual_ext = os.path.splitext(path)[1].lower().lstrip(".")
    if actual_ext in ("docx", "doc"):
        file_type = actual_ext

    if file_type not in ("doc", "docx"):
        raise HTTPException(status_code=415, detail="Preview only supported for DOCX files")

    cache_path = _docx_preview_cache_path(path)
    cached_html = _read_docx_preview_cache(cache_path)
    if cached_html is not None:
        return HTMLResponse(
            content=cached_html,
            headers={"Cache-Control": "public, max-age=3600"},
        )

    try:
        stat = os.stat(path)
        full_html = _build_cached_docx_html(path, stat.st_mtime_ns, stat.st_size)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {exc}") from exc
    try:
        tmp_path = f"{cache_path}.tmp"
        with open(tmp_path, "w", encoding="utf-8") as f:
            f.write(full_html)
        os.replace(tmp_path, cache_path)
    except OSError:
        pass

    return HTMLResponse(
        content=full_html,
        headers={"Cache-Control": "public, max-age=3600"},
    )


# ── Contract Versions ─────────────────────────────────────────────────────────

@router.get("/contracts/{contract_id}/versions", response_model=List[schemas.DocumentVersionSchema])
def read_versions(contract_id: int, db: Session = Depends(get_db)):
    return db.query(models.DocumentVersion).filter(
        models.DocumentVersion.contract_id == contract_id
    ).order_by(models.DocumentVersion.version_number.desc()).all()

@router.post("/contracts/{contract_id}/upload-version", response_model=schemas.ContractSchema)
async def upload_version(
    background_tasks: BackgroundTasks,
    contract_id: int,
    file: Optional[UploadFile] = File(None),
    drive_file_id: Optional[int] = Form(None),
    file_id: Optional[int] = Form(None),
    notes: Optional[str] = Form(None),
    label: Optional[str] = Form(None),
    db: Session = Depends(get_db),
):
    db_contract = db.query(models.Contract).filter(models.Contract.id == contract_id).first()
    if not db_contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    drive_file_id = file_id or drive_file_id

    last_version = db.query(models.DocumentVersion).filter(
        models.DocumentVersion.contract_id == contract_id
    ).order_by(models.DocumentVersion.version_number.desc()).first()

    next_num = (last_version.version_number + 1) if last_version else 1

    db_file: Optional[models.File] = None

    if drive_file_id is not None:
        db_file = db.query(models.File).filter(models.File.id == drive_file_id).first()
        if not db_file:
            raise HTTPException(status_code=404, detail="File not found")
    elif file:
        if not file.filename:
            raise HTTPException(status_code=400, detail="Uploaded file has no filename")
        file_type = file.filename.rsplit(".", 1)[-1].lower()
        file_path = os.path.join(UPLOAD_DIR, f"{contract_id}_v{next_num}_{int(datetime.now().timestamp())}.{file_type}")
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        db_file = _create_contract_upload_file(
            db,
            disk_path=file_path,
            original_filename=file.filename,
            file_type=file_type,
            content_type=file.content_type,
        )
    else:
        raise HTTPException(status_code=400, detail="Either file or drive_file_id is required")

    db_version = models.DocumentVersion(
        contract_id=contract_id,
        version_number=next_num,
        label=label or f"Version {next_num}",
        file_id=db_file.id,
        notes=notes,
        is_latest=False,
    )
    db.add(db_version)
    db.commit()
    db.refresh(db_version)

    _mark_version_latest(db, contract_id, db_version.id)

    db_contract.file_path = db_file.file_path
    db_contract.file_type = db_file.file_type or (
        db_file.original_filename.rsplit(".", 1)[-1].lower() if db_file.original_filename else None
    )
    _auto_milestone(db, contract_id, f"New Version Uploaded (v{next_num})",
                    f"Document version v{next_num} was uploaded." + (f" Notes: {notes}" if notes else ""))

    # Preserve lifecycle status (review, signing, etc.). New file ≠ new contract phase.
    db_contract.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(db_contract)

    existing = (
        db.query(models.DocumentChunk)
        .filter(models.DocumentChunk.file_id == db_file.id)
        .count()
    )
    if existing == 0:
        background_tasks.add_task(
            _run_chunk_and_compliance_background, contract_id, db_file.id
        )
    else:
        background_tasks.add_task(_run_post_chunk_tasks_background, contract_id, db_file.id)

    # Re-query with eager-loaded File relationships so Pydantic can read file_type / file_path.
    return (
        db.query(models.Contract)
        .options(
            selectinload(models.Contract.document_versions)
            .selectinload(models.DocumentVersion.file)
        )
        .filter(models.Contract.id == contract_id)
        .first()
    )


@router.delete("/contracts/{contract_id}/versions/{version_id}", response_model=schemas.ContractSchema)
def delete_document_version(contract_id: int, version_id: int, db: Session = Depends(get_db)):
    db_contract = db.query(models.Contract).filter(models.Contract.id == contract_id).first()
    if not db_contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    version = db.query(models.DocumentVersion).filter(
        models.DocumentVersion.id == version_id,
        models.DocumentVersion.contract_id == contract_id,
    ).first()
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")

    remaining = db.query(models.DocumentVersion).filter(
        models.DocumentVersion.contract_id == contract_id
    ).count()
    if remaining <= 1:
        raise HTTPException(status_code=400, detail="Cannot delete the only document version")

    signer_ids = [
        row[0]
        for row in db.query(models.VersionSigner.id).filter(
            models.VersionSigner.version_id == version_id
        ).all()
    ]
    sig_filters = [models.SignatureField.version_id == version_id]
    if signer_ids:
        sig_filters.append(models.SignatureField.version_signer_id.in_(signer_ids))
    db.query(models.SignatureField).filter(or_(*sig_filters)).delete(synchronize_session=False)

    was_latest = bool(version.is_latest)
    fid = version.file_id
    signed_file_path = version.signed_file_path

    db.delete(version)
    db.flush()

    if was_latest:
        new_latest = db.query(models.DocumentVersion).filter(
            models.DocumentVersion.contract_id == contract_id
        ).order_by(models.DocumentVersion.version_number.desc()).first()
        if new_latest:
            _mark_version_latest(db, contract_id, new_latest.id)
            db_contract.file_path = (new_latest.file.file_path if new_latest.file else None)
            db_contract.signed_file_path = new_latest.signed_file_path
            db_contract.file_type = (new_latest.file.file_type if new_latest.file else None)

    db_contract.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(db_contract)

    _safe_unlink(signed_file_path)

    # Contract-only ``File`` rows (not drive): remove disk + DB when no version references them.
    if fid:
        if (
            db.query(models.DocumentVersion)
            .filter(models.DocumentVersion.file_id == fid)
            .count()
            == 0
        ):
            frow = db.query(models.File).filter(models.File.id == fid).first()
            if frow and frow.folder_id is None:
                path = frow.file_path
                db.delete(frow)
                db.commit()
                _safe_unlink(path)

    return db_contract


# ── Version Signers ───────────────────────────────────────────────────────────

@router.get("/versions/{version_id}/signers", response_model=List[schemas.VersionSignerSchema])
def get_version_signers(version_id: int, db: Session = Depends(get_db)):
    version = db.query(models.DocumentVersion).filter(models.DocumentVersion.id == version_id).first()
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")
    return version.version_signers

@router.post("/versions/{version_id}/signers", response_model=schemas.VersionSignerSchema)
def add_signer_to_version(
    version_id: int,
    body: schemas.VersionSignerCreate,
    db: Session = Depends(get_db)
):
    version = db.query(models.DocumentVersion).filter(models.DocumentVersion.id == version_id).first()
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")

    master_signer = db.query(models.MasterSigner).filter(
        models.MasterSigner.id == body.master_signer_id
    ).first()
    if not master_signer:
        raise HTTPException(status_code=404, detail="Master signer not found")

    existing = db.query(models.VersionSigner).filter(
        models.VersionSigner.version_id == version_id,
        models.VersionSigner.master_signer_id == body.master_signer_id
    ).first()
    if existing:
        raise HTTPException(status_code=409, detail="This signer is already added to this version")

    db_vs = models.VersionSigner(
        version_id=version_id,
        master_signer_id=body.master_signer_id,
        signing_order=body.signing_order,
        token=str(uuid.uuid4()),
        status="pending",
    )
    db.add(db_vs)
    db.commit()
    db.refresh(db_vs)
    return db_vs

@router.delete("/version-signers/{vs_id}")
def remove_signer_from_version(vs_id: int, db: Session = Depends(get_db)):
    db_vs = db.query(models.VersionSigner).filter(models.VersionSigner.id == vs_id).first()
    if not db_vs:
        raise HTTPException(status_code=404, detail="Version signer not found")
    if db_vs.status == "signed":
        raise HTTPException(status_code=400, detail="Cannot remove a signer who has already signed")
    # Also delete their signature fields
    db.query(models.SignatureField).filter(
        models.SignatureField.version_signer_id == vs_id
    ).delete()
    db.delete(db_vs)
    db.commit()
    return {"detail": "Signer removed from version"}


# ── Signature Fields (version-linked) ────────────────────────────────────────

@router.post("/versions/{version_id}/signature-fields", response_model=schemas.SignatureFieldSchema)
def add_signature_field_to_version(
    version_id: int,
    field: schemas.SignatureFieldBase,
    db: Session = Depends(get_db)
):
    version = db.query(models.DocumentVersion).filter(models.DocumentVersion.id == version_id).first()
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")

    data = field.model_dump()
    data["version_id"] = version_id
    # Map legacy signer_id → version_signer_id if provided
    if data.get("version_signer_id"):
        data["signer_id"] = None
    db_field = models.SignatureField(**data)
    db.add(db_field)
    db.commit()
    db.refresh(db_field)
    return db_field

@router.get("/versions/{version_id}/signature-fields", response_model=List[schemas.SignatureFieldSchema])
def get_version_signature_fields(version_id: int, db: Session = Depends(get_db)):
    return db.query(models.SignatureField).filter(
        models.SignatureField.version_id == version_id
    ).all()

# Legacy contract-level field endpoint kept for backward compat
@router.post("/contracts/{contract_id}/signature-fields", response_model=schemas.SignatureFieldSchema)
def add_signature_field(contract_id: int, field: schemas.SignatureFieldBase, db: Session = Depends(get_db)):
    db_contract = db.query(models.Contract).filter(models.Contract.id == contract_id).first()
    if not db_contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    data = field.model_dump()
    data["contract_id"] = contract_id
    db_field = models.SignatureField(**data)
    db.add(db_field)
    db.commit()
    db.refresh(db_field)
    return db_field

@router.patch("/signature-fields/{field_id}", response_model=schemas.SignatureFieldSchema)
def update_signature_field(field_id: int, field_update: schemas.SignatureFieldUpdate, db: Session = Depends(get_db)):
    db_field = db.query(models.SignatureField).filter(models.SignatureField.id == field_id).first()
    if not db_field:
        raise HTTPException(status_code=404, detail="Signature field not found")
    for key, value in field_update.model_dump(exclude_unset=True).items():
        setattr(db_field, key, value)
    db.commit()
    db.refresh(db_field)
    return db_field

@router.delete("/signature-fields/{field_id}")
def delete_signature_field(field_id: int, db: Session = Depends(get_db)):
    db_field = db.query(models.SignatureField).filter(models.SignatureField.id == field_id).first()
    if not db_field:
        raise HTTPException(status_code=404, detail="Signature field not found")
    db.delete(db_field)
    db.commit()
    return {"detail": "Signature field deleted"}


# ── Send Invitations (version-aware) ─────────────────────────────────────────

@router.post("/versions/{version_id}/send-invitations")
def send_version_invitations(version_id: int, db: Session = Depends(get_db)):
    version = db.query(models.DocumentVersion).filter(models.DocumentVersion.id == version_id).first()
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")

    version_signers = db.query(models.VersionSigner).filter(
        models.VersionSigner.version_id == version_id,
        models.VersionSigner.status == "pending"
    ).all()

    for vs in version_signers:
        vs.status = "invited"
        vs.invited_at = datetime.utcnow()

    # Update contract status
    db_contract = db.query(models.Contract).filter(
        models.Contract.id == version.contract_id
    ).first()
    if db_contract:
        db_contract.status = models.ContractStatus.SIGNING
        # Auto milestone: sent for signing
        signer_names = ", ".join(vs.master_signer.name for vs in version_signers if vs.master_signer)
        _auto_milestone(db, version.contract_id, "Sent for Signing",
                        f"Version v{version.version_number} sent to {len(version_signers)} signer(s): {signer_names}.",
                        status="completed")

    db.commit()
    return {
        "message": f"Invitations sent to {len(version_signers)} signers",
        "status": "signing",
        "signing_links": [
            {"signer": vs.master_signer.name, "token": vs.token}
            for vs in version_signers
        ]
    }

# Legacy contract-level send-invitations
@router.post("/contracts/{contract_id}/send-invitations")
def send_invitations(contract_id: int, db: Session = Depends(get_db)):
    db_contract = db.query(models.Contract).filter(models.Contract.id == contract_id).first()
    if not db_contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    db_contract.status = models.ContractStatus.SIGNING
    db.commit()
    signers = db.query(models.Signer).filter(models.Signer.contract_id == contract_id).all()
    return {"message": f"Invitations sent to {len(signers)} signers", "status": "signing"}


# ── Token-based Signing ───────────────────────────────────────────────────────

@router.get("/sign/{token}", response_model=schemas.SigningContextSchema)
def get_signing_context(token: str, db: Session = Depends(get_db)):
    vs = db.query(models.VersionSigner).filter(models.VersionSigner.token == token).first()
    if not vs:
        raise HTTPException(status_code=404, detail="Invalid or expired signing link")
    if vs.status == "signed":
        raise HTTPException(status_code=400, detail="This document has already been signed")

    version = vs.version
    contract = version.contract_parent
    fields = db.query(models.SignatureField).filter(
        models.SignatureField.version_id == version.id
    ).all()

    return schemas.SigningContextSchema(
        token=token,
        version_signer=vs,
        contract_id=contract.id,
        contract_title=contract.title,
        version_id=version.id,
        version_number=version.version_number,
        file_type=(version.file.file_type if version.file else None) or "pdf",
        all_version_signers=version.version_signers,
        signature_fields=fields,
    )

@router.post("/sign/{token}/complete")
def complete_signing(token: str, db: Session = Depends(get_db)):
    vs = db.query(models.VersionSigner).filter(models.VersionSigner.token == token).first()
    if not vs:
        raise HTTPException(status_code=404, detail="Invalid signing link")
    if vs.status == "signed":
        raise HTTPException(status_code=400, detail="Already signed")

    vs.status = "signed"
    vs.signed_at = datetime.utcnow()

    # Mark this signer's fields as signed
    db.query(models.SignatureField).filter(
        models.SignatureField.version_signer_id == vs.id
    ).update({"is_signed": True})

    db.commit()

    # Check if all signers for this version have signed
    version = vs.version
    all_signers = db.query(models.VersionSigner).filter(
        models.VersionSigner.version_id == version.id
    ).all()
    all_signed = all(s.status == "signed" for s in all_signers)

    contract = version.contract_parent
    signer_name = vs.master_signer.name if vs.master_signer else "Unknown"

    # Auto milestone: individual signer signed
    _auto_milestone(db, contract.id, f"Signed by {signer_name}",
                    f"{signer_name} signed version v{version.version_number}.")

    if all_signed:
        contract.status = models.ContractStatus.ACTIVE
        # Auto milestone: fully signed / contract activated
        _auto_milestone(db, contract.id, "Contract Fully Signed",
                        f"All signers completed signing version v{version.version_number}. Contract is now active.")
        db.commit()
        try:
            ft = ((version.file.file_type if version.file else None) or "pdf").lower()
            if ft in ("docx", "doc"):
                generate_signed_docx_v2(version, db)
            else:
                generate_signed_pdf_v2(version, db)
            db.commit()
        except Exception:
            logger.exception("Error generating signed document.")
    else:
        db.commit()

    return {"message": "Successfully signed", "all_signed": all_signed}


# ── Milestones ────────────────────────────────────────────────────────────────

@router.get("/contracts/{contract_id}/milestones", response_model=List[schemas.MilestoneSchema])
def read_milestones(contract_id: int, db: Session = Depends(get_db)):
    return db.query(models.Milestone).filter(models.Milestone.contract_id == contract_id).all()

@router.post("/contracts/{contract_id}/milestones", response_model=schemas.MilestoneSchema)
def create_milestone(contract_id: int, milestone: schemas.MilestoneBase, db: Session = Depends(get_db)):
    db_contract = db.query(models.Contract).filter(models.Contract.id == contract_id).first()
    if not db_contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    db_milestone = models.Milestone(**milestone.model_dump(), contract_id=contract_id)
    db.add(db_milestone)
    db.commit()
    db.refresh(db_milestone)
    return db_milestone

@router.patch("/milestones/{milestone_id}", response_model=schemas.MilestoneSchema)
def update_milestone(milestone_id: int, milestone_update: schemas.MilestoneUpdate, db: Session = Depends(get_db)):
    db_milestone = db.query(models.Milestone).filter(models.Milestone.id == milestone_id).first()
    if not db_milestone:
        raise HTTPException(status_code=404, detail="Milestone not found")
    db_milestone.status = milestone_update.status
    db.commit()
    db.refresh(db_milestone)
    return db_milestone


# ── Compliance ────────────────────────────────────────────────────────────────

@router.post(
    "/contracts/{contract_id}/compliance/recompute",
    response_model=schemas.ComplianceRecomputeResponse,
)
async def recompute_compliance_checks(
    contract_id: int,
    body: schemas.ComplianceRecomputeRequest,
    db: Session = Depends(get_db),
):
    """
    Re-run structured LLM compliance for the given ``version_id`` (must belong to the contract).

    Loads the latest ``document_chunks`` for that version's ``file_id``. Existing automated compliance \
    rows for **this document version only** are replaced after a successful LLM response (one transaction).
    """
    from app.services.compliance_check import (
        generate_compliance_checks_from_file,
        record_compliance_run_milestone,
        replace_compliance_records_atomic,
    )

    db_contract = db.query(models.Contract).filter(models.Contract.id == contract_id).first()
    if not db_contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    ver = (
        db.query(models.DocumentVersion)
        .filter(
            models.DocumentVersion.id == body.version_id,
            models.DocumentVersion.contract_id == contract_id,
        )
        .first()
    )
    if not ver:
        raise HTTPException(status_code=404, detail="Document version not found for this contract")
    if not ver.file_id:
        raise HTTPException(status_code=400, detail="Version has no linked file")

    try:
        response = await generate_compliance_checks_from_file(
            db,
            ver.file_id,
            contract_id=contract_id,
            document_version_id=ver.id,
        )
    except ValueError as e:
        if str(e) == "no_chunks":
            raise HTTPException(
                status_code=400,
                detail=(
                    "No document chunks for this file yet. Wait for indexing to finish after upload, "
                    "or open the document preview once to trigger processing."
                ),
            ) from e
        raise
    except Exception as e:
        logger.exception("Compliance LLM failed contract_id=%s version_id=%s", contract_id, body.version_id)
        raise HTTPException(
            status_code=502,
            detail=f"Compliance generation failed: {e!s}",
        ) from e

    try:
        count = replace_compliance_records_atomic(db, contract_id, response, ver.id)
    except Exception as e:
        logger.exception("Compliance DB replace failed contract_id=%s", contract_id)
        raise HTTPException(
            status_code=500,
            detail="Failed to save compliance records; existing data was left unchanged.",
        ) from e

    try:
        record_compliance_run_milestone(
            db,
            contract_id,
            ver.file_id,
            count=count,
            trigger="Manual compliance re-run",
            version_id=ver.id,
        )
    except Exception:
        logger.exception("Compliance milestone failed contract_id=%s", contract_id)

    return schemas.ComplianceRecomputeResponse(
        contract_id=contract_id,
        version_id=ver.id,
        file_id=ver.file_id,
        count=count,
    )


@router.get("/contracts/{contract_id}/compliance", response_model=List[schemas.ComplianceRecordSchema])
def read_compliance_records(
    contract_id: int,
    record_type: Optional[str] = None,
    version_id: Optional[int] = Query(None),
    db: Session = Depends(get_db),
):
    query = db.query(models.ComplianceRecord).filter(models.ComplianceRecord.contract_id == contract_id)
    if record_type:
        query = query.filter(models.ComplianceRecord.record_type == record_type)
    if version_id is not None:
        query = query.filter(models.ComplianceRecord.document_version_id == version_id)
    return query.all()

@router.post("/contracts/{contract_id}/compliance", response_model=schemas.ComplianceRecordSchema)
def create_compliance_record(contract_id: int, record: schemas.ComplianceRecordBase, db: Session = Depends(get_db)):
    db_contract = db.query(models.Contract).filter(models.Contract.id == contract_id).first()
    if not db_contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    db_record = models.ComplianceRecord(**record.model_dump(), contract_id=contract_id)
    db.add(db_record)
    db.commit()
    db.refresh(db_record)
    rt = (record.record_type or "compliance").lower()
    if rt == "compliance":
        _auto_milestone(
            db,
            contract_id,
            "Compliance entry added",
            f"Manual entry: {record.check_name} ({record.status}). "
            f"{(record.findings or '')[:900]}",
        )
        db.commit()
    return db_record

@router.get("/complaints", response_model=List[schemas.ComplianceRecordSchema])
def get_all_complaints(db: Session = Depends(get_db)):
    return db.query(models.ComplianceRecord).all()


# ── Legacy Signers (kept for backward compat) ─────────────────────────────────

@router.post("/contracts/{contract_id}/signers", response_model=schemas.SignerSchema)
def add_signer(contract_id: int, signer: schemas.SignerBase, db: Session = Depends(get_db)):
    db_contract = db.query(models.Contract).filter(models.Contract.id == contract_id).first()
    if not db_contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    db_signer = models.Signer(contract_id=contract_id, **signer.model_dump())
    db.add(db_signer)
    db.commit()
    db.refresh(db_signer)
    return db_signer

@router.post("/signers/{signer_id}/sign")
def sign_contract_legacy(signer_id: int, db: Session = Depends(get_db)):
    db_signer = db.query(models.Signer).filter(models.Signer.id == signer_id).first()
    if not db_signer:
        raise HTTPException(status_code=404, detail="Signer not found")
    db_signer.status = "signed"
    db_signer.signed_at = datetime.utcnow()
    db.query(models.SignatureField).filter(
        models.SignatureField.signer_id == signer_id
    ).update({"is_signed": True})
    db.commit()

    db_contract = db.query(models.Contract).filter(models.Contract.id == db_signer.contract_id).first()
    all_signers = db.query(models.Signer).filter(models.Signer.contract_id == db_signer.contract_id).all()
    if all(s.status == "signed" for s in all_signers):
        db_contract.status = models.ContractStatus.ACTIVE
    db.commit()
    try:
        generate_signed_pdf(db_contract, db)
        db.commit()
    except Exception:
        logger.exception("Error generating signed PDF.")
    return {"message": "Successfully signed"}


# ── PDF Generation ────────────────────────────────────────────────────────────

def _draw_fields_on_canvas(c, fields, get_signer_name_fn, page_height):
    for field in fields:
        if not field.is_signed or not field.value:
            continue
        scale = field.scale or 1.5
        x_pt = field.x_pos / scale
        field_height = field.height or 50
        field_width = field.width or 150
        y_pt = page_height - (field.y_pos / scale) - (field_height / scale)

        c.setDash(3, 3)
        c.setStrokeColorRGB(0.7, 0.7, 0.7)
        c.rect(x_pt, y_pt, field_width / scale, field_height / scale, stroke=1, fill=0)

        signer_name = get_signer_name_fn(field)
        if signer_name:
            c.setFont("Helvetica", 8)
            c.setFillColorRGB(0.3, 0.3, 0.3)
            c.drawString(x_pt, y_pt + (field_height / scale) + 2, signer_name)

        c.setDash([])
        c.setStrokeColorRGB(0, 0, 0)
        c.setFillColorRGB(0, 0, 0)

        if field.field_type == "signature" and field.value.startswith("data:image"):
            try:
                image_bytes = base64.b64decode(field.value.split(",")[1])
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as img_temp:
                    img_temp.write(image_bytes)
                    img_path = img_temp.name
                c.drawImage(img_path, x_pt, y_pt, width=field_width / scale, height=field_height / scale, mask="auto")
                os.unlink(img_path)
            except Exception:
                logger.exception("Error processing signature image.")
        elif field.field_type == "text":
            text_width = c.stringWidth(field.value, "Helvetica", 10)
            c.setFont("Helvetica", 10)
            c.drawString(x_pt + (field_width / scale - text_width) / 2, y_pt + (field_height / scale - 10) / 2 + 2, field.value)
        elif field.field_type == "date":
            date_str = f"Date: {field.value}"
            text_width = c.stringWidth(date_str, "Helvetica", 10)
            c.setFont("Helvetica", 10)
            c.drawString(x_pt + (field_width / scale - text_width) / 2, y_pt + (field_height / scale - 10) / 2 + 2, date_str)

def generate_signed_docx_v2(version: models.DocumentVersion, db: Session):
    """Append a digital-signature certificate page to the DOCX and save as signed copy."""
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from collections import defaultdict

    _vfp = version.file.file_path if version.file else None
    if not _vfp or not os.path.exists(_vfp):
        return
    if ((version.file.file_type if version.file else None) or "").lower() not in ("docx", "doc"):
        return

    fields = db.query(models.SignatureField).filter(
        models.SignatureField.version_id == version.id,
        models.SignatureField.is_signed == True,
    ).all()

    if not fields:
        return

    signer_fields: dict = defaultdict(list)
    for f in fields:
        signer_fields[f.version_signer_id].append(f)

    doc = DocxDocument(_vfp)
    doc.add_page_break()

    h = doc.add_heading("Digital Signature Certificate", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ts = doc.add_paragraph(f"Signed on: {datetime.utcnow().strftime('%B %d, %Y at %H:%M UTC')}")
    ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    for vs_id, vs_fields in signer_fields.items():
        vs = db.query(models.VersionSigner).filter(models.VersionSigner.id == vs_id).first()
        name    = (vs.master_signer.name  if vs and vs.master_signer else "Unknown")
        email   = (vs.master_signer.email if vs and vs.master_signer else "")
        signed  = (vs.signed_at.strftime("%B %d, %Y %H:%M UTC") if vs and vs.signed_at else "")

        sh = doc.add_heading(name, level=2)
        if email:   doc.add_paragraph(f"Email:  {email}")
        if signed:  doc.add_paragraph(f"Signed: {signed}")

        for field in sorted(vs_fields, key=lambda x: (x.page_number or 0, x.y_pos or 0)):
            if field.field_type == "signature" and field.value and field.value.startswith("data:image"):
                try:
                    img_bytes = base64.b64decode(field.value.split(",")[1])
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                        tmp.write(img_bytes)
                        img_path = tmp.name
                    p = doc.add_paragraph("Signature: ")
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(2.2))
                    os.unlink(img_path)
                except Exception:
                    logger.exception("Signature embed error.")
                    doc.add_paragraph("[Signature]")
            elif field.field_type == "initials" and field.value:
                doc.add_paragraph(f"Initials: {field.value}")
            elif field.field_type == "date" and field.value:
                doc.add_paragraph(f"Date: {field.value}")
            elif field.field_type == "text" and field.value:
                doc.add_paragraph(f"Text: {field.value}")

        doc.add_paragraph("")   # spacer between signers

    signed_path = os.path.join(
        UPLOAD_DIR, f"signed_v{version.id}_{int(datetime.now().timestamp())}.docx"
    )
    doc.save(signed_path)
    version.signed_file_path = signed_path


def generate_signed_pdf_v2(version: models.DocumentVersion, db: Session):
    """Generate a signed PDF for a version using version_signer-linked fields."""
    _vfp = version.file.file_path if version.file else None
    if not _vfp or not os.path.exists(_vfp) or (version.file.file_type if version.file else None) != "pdf":
        return

    fields = db.query(models.SignatureField).filter(
        models.SignatureField.version_id == version.id
    ).all()

    pages_fields = {}
    for field in fields:
        if not field.is_signed or not field.value:
            continue
        idx = field.page_number if field.page_number is not None else 0
        pages_fields.setdefault(idx, []).append(field)

    def get_name(field):
        if field.version_signer and field.version_signer.master_signer:
            return field.version_signer.master_signer.name
        return ""

    reader = PdfReader(_vfp)
    writer = PdfWriter()

    for i in range(len(reader.pages)):
        original_page = reader.pages[i]
        if i in pages_fields:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as ot:
                overlay_path = ot.name
            mb = original_page.mediabox
            page_width, page_height = float(mb.width), float(mb.height)
            c = canvas.Canvas(overlay_path, pagesize=(page_width, page_height))
            _draw_fields_on_canvas(c, pages_fields[i], get_name, page_height)
            c.save()
            overlay_page = PdfReader(overlay_path).pages[0]
            original_page.merge_page(overlay_page)
            os.unlink(overlay_path)
        writer.add_page(original_page)

    signed_path = os.path.join(UPLOAD_DIR, f"signed_v{version.id}_{int(datetime.now().timestamp())}.pdf")
    with open(signed_path, "wb") as f:
        writer.write(f)
    version.signed_file_path = signed_path

def generate_signed_pdf(contract: models.Contract, db: Session):
    """Legacy: generate signed PDF using contract-level fields."""
    if not contract.file_path or not os.path.exists(contract.file_path) or contract.file_type != "pdf":
        return

    fields = db.query(models.SignatureField).filter(
        models.SignatureField.contract_id == contract.id
    ).all()

    pages_fields = {}
    for field in fields:
        if not field.is_signed or not field.value:
            continue
        idx = field.page_number if field.page_number is not None else 0
        pages_fields.setdefault(idx, []).append(field)

    def get_name(field):
        return field.signer.name if field.signer else ""

    reader = PdfReader(contract.file_path)
    writer = PdfWriter()

    for i in range(len(reader.pages)):
        original_page = reader.pages[i]
        if i in pages_fields:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as ot:
                overlay_path = ot.name
            mb = original_page.mediabox
            page_height = float(mb.height)
            c = canvas.Canvas(overlay_path, pagesize=(float(mb.width), page_height))
            _draw_fields_on_canvas(c, pages_fields[i], get_name, page_height)
            c.save()
            overlay_page = PdfReader(overlay_path).pages[0]
            original_page.merge_page(overlay_page)
            os.unlink(overlay_path)
        writer.add_page(original_page)

    signed_path = os.path.join(UPLOAD_DIR, f"signed_{contract.id}_{int(datetime.now().timestamp())}.pdf")
    with open(signed_path, "wb") as f:
        writer.write(f)
    contract.signed_file_path = signed_path


# ── LLM Result (notebook) ─────────────────────────────────────────────────────

def _llm_result_json_path() -> Path:
    """Resolve path to llm_result.json; read from disk on every request."""
    env_path = os.environ.get("LLM_RESULT_JSON_PATH")
    if env_path:
        return Path(env_path).resolve()
    # Default: clm/notebook/llm_result.json relative to project root
    project_root = Path(__file__).resolve().parents[4]
    return project_root / "clm" / "notebook" / "llm_result.json"


# ── Review Items ─────────────────────────────────────────────────────────────

@router.get("/contracts/{contract_id}/review-items", response_model=List[schemas.ReviewItemSchema])
def list_review_items(contract_id: int, db: Session = Depends(get_db)):
    return (
        db.query(models.ReviewItem)
        .filter(models.ReviewItem.contract_id == contract_id)
        .order_by(models.ReviewItem.created_at.desc())
        .all()
    )


@router.post("/contracts/{contract_id}/review-items", response_model=schemas.ReviewItemSchema, status_code=201)
def create_review_item(contract_id: int, body: schemas.ReviewItemCreate, db: Session = Depends(get_db)):
    contract = db.query(models.Contract).filter(models.Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    item = models.ReviewItem(
        contract_id=contract_id,
        title=body.title,
        content=body.content,
        source_query=body.source_query,
        item_type=body.item_type or "finding",
        severity=body.severity,
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return item


@router.delete("/review-items/{item_id}", status_code=204)
def delete_review_item(item_id: int, db: Session = Depends(get_db)):
    item = db.query(models.ReviewItem).filter(models.ReviewItem.id == item_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="Review item not found")
    db.delete(item)
    db.commit()


@router.get("/llm-result")
def get_llm_result():
    """Read llm_result.json from disk on every request; no caching."""
    path = _llm_result_json_path()
    if not path.exists():
        raise HTTPException(
            status_code=404,
            detail=f"llm_result.json not found at {path}. Set LLM_RESULT_JSON_PATH to override.",
        )
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=500, detail=f"Invalid JSON in llm_result.json: {e}")


# ── Contract Scoring (agent workflow) ──────────────────────────────────────────


def _scoring_version_for_contract(
    db: Session,
    contract_id: int,
    *,
    version_id: Optional[int] = None,
    file_id: Optional[int] = None,
) -> Optional[models.DocumentVersion]:
    """Pick the document version used for scoring: explicit version, file, or contract latest."""
    if version_id is not None:
        return (
            db.query(models.DocumentVersion)
            .filter(
                models.DocumentVersion.id == version_id,
                models.DocumentVersion.contract_id == contract_id,
            )
            .first()
        )
    if file_id is not None:
        return (
            db.query(models.DocumentVersion)
            .filter(
                models.DocumentVersion.contract_id == contract_id,
                models.DocumentVersion.file_id == file_id,
            )
            .order_by(models.DocumentVersion.version_number.desc())
            .first()
        )
    return (
        db.query(models.DocumentVersion)
        .filter(
            models.DocumentVersion.contract_id == contract_id,
            models.DocumentVersion.is_latest == True,
        )
        .first()
    )


def _file_belongs_to_contract(db: Session, contract_id: int, file_id: int) -> bool:
    return (
        db.query(models.DocumentVersion)
        .filter(
            models.DocumentVersion.contract_id == contract_id,
            models.DocumentVersion.file_id == file_id,
        )
        .first()
        is not None
    )


@router.get(
    "/contracts/{contract_id}/scoring",
    response_model=schemas.ScoringResultReadSchema,
)
def get_contract_scoring(
    contract_id: int,
    db: Session = Depends(get_db),
    version_id: Optional[int] = Query(None, description="Return score for this document version (recommended)"),
    file_id: Optional[int] = Query(None, description="Return score for this file id (alternate to version_id)"),
):
    """Return the latest scoring result for the contract, optionally filtered by version or file."""
    contract = db.query(models.Contract).filter(models.Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    base = db.query(models.ScoringResult).filter(models.ScoringResult.contract_id == contract_id)

    if version_id is not None:
        resolved = _scoring_version_for_contract(db, contract_id, version_id=version_id, file_id=None)
        if not resolved:
            raise HTTPException(status_code=404, detail="Document version not found for this contract")
        row = None
        if resolved.file_id:
            row = (
                base.filter(models.ScoringResult.file_id == resolved.file_id)
                .order_by(models.ScoringResult.created_at.desc())
                .first()
            )
        if row is None:
            row = (
                base.filter(models.ScoringResult.document_version_id == resolved.id)
                .order_by(models.ScoringResult.created_at.desc())
                .first()
            )
        if not row:
            return schemas.ScoringResultReadSchema(
                contract_id=contract_id,
                document_version_id=resolved.id,
                file_id=resolved.file_id,
                result_json=None,
            )
        return row

    if file_id is not None:
        if not _file_belongs_to_contract(db, contract_id, file_id):
            raise HTTPException(status_code=404, detail="File is not associated with this contract")
        row = (
            base.filter(models.ScoringResult.file_id == file_id)
            .order_by(models.ScoringResult.created_at.desc())
            .first()
        )
        if row is None:
            resolved = _scoring_version_for_contract(db, contract_id, file_id=file_id)
            return schemas.ScoringResultReadSchema(
                contract_id=contract_id,
                document_version_id=resolved.id if resolved else None,
                file_id=file_id,
                result_json=None,
            )
        return row

    row = base.order_by(models.ScoringResult.created_at.desc()).first()
    if not row:
        return schemas.ScoringResultReadSchema(contract_id=contract_id, result_json=None)
    return row


@router.post(
    "/contracts/{contract_id}/scoring",
    response_model=schemas.ScoringResultSchema,
    status_code=201,
)
async def trigger_contract_scoring(
    contract_id: int,
    db: Session = Depends(get_db),
    version_id: Optional[int] = Query(None, description="Score this document version (recommended)"),
    file_id: Optional[int] = Query(None, description="Score the version that uses this file id"),
):
    """Run the multi-agent scoring workflow and store the result for the selected file/version."""
    import asyncio

    contract = db.query(models.Contract).filter(models.Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    if version_id is not None and file_id is not None:
        v_chk = _scoring_version_for_contract(db, contract_id, version_id=version_id, file_id=None)
        if v_chk and v_chk.file_id != file_id:
            raise HTTPException(
                status_code=400,
                detail="version_id and file_id refer to different documents",
            )

    version = _scoring_version_for_contract(db, contract_id, version_id=version_id, file_id=file_id)
    if not version:
        raise HTTPException(status_code=404, detail="No document version found for scoring")
    if not version.file_id:
        raise HTTPException(
            status_code=400,
            detail="Selected version has no file; upload a document first.",
        )

    from app.services.agent_workflow import compile_graph, load_contract_guidelines

    try:
        initial_state = load_contract_guidelines(contract_id, db)
    except ValueError as exc:
        print(f"Error loading contract guidelines: {exc}")
        raise HTTPException(status_code=404, detail=str(exc))

    if not initial_state.get("contract_text"):
        chunks = (
            db.query(models.DocumentChunk)
            .filter(models.DocumentChunk.file_id == version.file_id)
            .order_by(models.DocumentChunk.chunk_index)
            .all()
        )
        if chunks:
            initial_state["contract_text"] = "\n\n".join(ch.content for ch in chunks)
            logger.info(
                "Built contract_text from %d chunks for version %d (file_id=%s)",
                len(chunks),
                version.id,
                version.file_id,
            )

    if not initial_state.get("contract_text"):
        raise HTTPException(
            status_code=400,
            detail="No contract text available. Ensure the document has been chunked first.",
        )

    graph = compile_graph()
    result_state = await asyncio.to_thread(graph.invoke, initial_state)

    result_json = {
        "contract_text": result_state.get("contract_text", ""),
        "contract_metadata": result_state.get("contract_metadata", {}),
        "classifications": result_state.get("classifications", {}),
        "validation_results": result_state.get("validation_results", []),
        "finance_analysis": result_state.get("finance_analysis", {}),
        "cross_validation_result": result_state.get("cross_validation_result", {}),
    }

    row = models.ScoringResult(
        contract_id=contract_id,
        document_version_id=version.id,
        file_id=version.file_id,
        result_json=result_json,
    )
    db.add(row)
    _auto_milestone(
        db, contract_id,
        title="Scoring completed",
        description=f"Multi-agent scoring workflow completed for document version v{version.version_number}.",
    )
    db.commit()
    db.refresh(row)

    logger.info(
        "Scoring completed for contract %d, result id=%d, file_id=%s",
        contract_id,
        row.id,
        row.file_id,
    )
    return row


@router.get("/contracts/{contract_id}/versions/{version_id}/knowledge-graph")
def get_contract_version_knowledge_graph(
    contract_id: int,
    version_id: int,
    db: Session = Depends(get_db),
):
    """
    Return stored LangChain graph JSON (array of GraphDocument dicts) for ``parseGraphData`` on the client.
    One row per ``document_version_id``.
    """
    row = (
        db.query(models.VersionKnowledgeGraph)
        .filter(
            models.VersionKnowledgeGraph.contract_id == contract_id,
            models.VersionKnowledgeGraph.document_version_id == version_id,
        )
        .first()
    )
    if not row:
        raise HTTPException(
            status_code=404,
            detail="No knowledge graph for this version. Run POST …/knowledge-graph/build first.",
        )
    return row.graph_documents_json


@router.post(
    "/contracts/{contract_id}/versions/{version_id}/knowledge-graph/build",
    response_model=schemas.VersionKnowledgeGraphBuildResponse,
    status_code=201,
)
async def build_contract_version_knowledge_graph(
    contract_id: int,
    version_id: int,
    db: Session = Depends(get_db),
):
    """Chunk text → LangChain Documents → ``LLMGraphTransformer`` → store ``graph_documents_json``."""
    import asyncio

    from app.services import graph_builder as graph_builder_svc

    try:
        full_text, meta = graph_builder_svc.load_version_text_for_graph(db, contract_id, version_id)
    except ValueError as exc:
        msg = str(exc)
        if "not found" in msg.lower():
            raise HTTPException(status_code=404, detail=msg) from exc
        raise HTTPException(status_code=400, detail=msg) from exc

    if not (full_text or "").strip():
        raise HTTPException(status_code=400, detail="Document text is empty")

    try:
        graph_docs = await asyncio.to_thread(
            graph_builder_svc.transform_text_to_graph_list,
            full_text,
            meta,
        )
    except Exception as exc:
        logger.exception(
            "Knowledge graph LLM failed contract_id=%s version_id=%s",
            contract_id,
            version_id,
        )
        raise HTTPException(
            status_code=502,
            detail=f"Knowledge graph extraction failed: {exc!s}",
        ) from exc

    row = graph_builder_svc.persist_version_knowledge_graph(db, contract_id, version_id, graph_docs)
    _auto_milestone(
        db, contract_id,
        title="Knowledge graph built",
        description=f"Knowledge graph built for document version {version_id}.",
    )
    db.commit()
    return schemas.VersionKnowledgeGraphBuildResponse(
        contract_id=row.contract_id,
        document_version_id=row.document_version_id,
        graph_documents=row.graph_documents_json,
        created_at=row.created_at,
        updated_at=row.updated_at,
    )


# ── Document Drive repositories ────────────────────────────────────────────────

def _get_drive_or_404(db: Session, drive_id: int):
    drive = (
        db.query(models.DocumentDrive)
        .options(selectinload(models.DocumentDrive.folders).selectinload(models.DocumentDriveFolder.files))
        .filter(models.DocumentDrive.id == drive_id)
        .first()
    )
    if not drive:
        raise HTTPException(status_code=404, detail="Document drive not found")
    return drive


def _get_folder_or_404(db: Session, folder_id: int):
    folder = (
        db.query(models.DocumentDriveFolder)
        .options(selectinload(models.DocumentDriveFolder.files))
        .filter(models.DocumentDriveFolder.id == folder_id)
        .first()
    )
    if not folder:
        raise HTTPException(status_code=404, detail="Folder not found")
    return folder


@router.get("/document-drives", response_model=List[schemas.DocumentDriveSchema])
def list_document_drives(db: Session = Depends(get_db)):
    drives = (
        db.query(models.DocumentDrive)
        .options(selectinload(models.DocumentDrive.folders).selectinload(models.DocumentDriveFolder.files))
        .order_by(models.DocumentDrive.created_at.desc())
        .all()
    )
    return drives


@router.post("/document-drives", response_model=schemas.DocumentDriveSchema)
def create_document_drive(payload: schemas.DocumentDriveCreate, db: Session = Depends(get_db)):
    name = (payload.name or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="Drive name is required")
    drive = models.DocumentDrive(
        name=name,
        description=payload.description,
    )
    db.add(drive)
    db.commit()
    db.refresh(drive)
    return drive


@router.get("/document-drives/{drive_id}", response_model=schemas.DocumentDriveSchema)
def get_document_drive(drive_id: int, db: Session = Depends(get_db)):
    drive = _get_drive_or_404(db, drive_id)
    return drive


@router.patch("/document-drives/{drive_id}", response_model=schemas.DocumentDriveSchema)
def update_document_drive(drive_id: int, payload: schemas.DocumentDriveCreate, db: Session = Depends(get_db)):
    drive = _get_drive_or_404(db, drive_id)
    if payload.name is not None:
        new_name = payload.name.strip()
        if not new_name:
            raise HTTPException(status_code=400, detail="Drive name is required")
        drive.name = new_name
    if payload.description is not None:
        drive.description = payload.description
    db.commit()
    db.refresh(drive)
    return drive


@router.delete("/document-drives/{drive_id}")
def delete_document_drive(drive_id: int, db: Session = Depends(get_db)):
    drive = _get_drive_or_404(db, drive_id)
    db.delete(drive)
    db.commit()
    return {"detail": "Document drive deleted"}


def _eager_folder_options():
    return selectinload(models.DocumentDriveFolder.files), selectinload(models.DocumentDriveFolder.children).selectinload(models.DocumentDriveFolder.files)


def _load_folder_tree(db: Session, drive_id: int):
    """Return all folders for a drive, structured as a flat list (frontend builds tree via parent_id)."""
    return (
        db.query(models.DocumentDriveFolder)
        .options(selectinload(models.DocumentDriveFolder.files))
        .filter(models.DocumentDriveFolder.drive_id == drive_id)
        .order_by(models.DocumentDriveFolder.created_at.asc())
        .all()
    )


@router.get("/document-drives/{drive_id}/folders", response_model=List[schemas.DocumentDriveFolderSchema])
def list_drive_folders(drive_id: int, db: Session = Depends(get_db)):
    _get_drive_or_404(db, drive_id)
    return _load_folder_tree(db, drive_id)


@router.post("/document-drives/{drive_id}/folders", response_model=schemas.DocumentDriveFolderSchema)
def create_drive_folder(
    drive_id: int,
    payload: schemas.DocumentDriveFolderCreate,
    db: Session = Depends(get_db)
):
    _get_drive_or_404(db, drive_id)
    name = (payload.name or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="Folder name is required")
    if payload.parent_id is not None:
        parent = db.query(models.DocumentDriveFolder).filter(
            models.DocumentDriveFolder.id == payload.parent_id,
            models.DocumentDriveFolder.drive_id == drive_id,
        ).first()
        if not parent:
            raise HTTPException(status_code=404, detail="Parent folder not found")
    folder = models.DocumentDriveFolder(
        drive_id=drive_id,
        parent_id=payload.parent_id,
        name=name,
        description=payload.description,
    )
    db.add(folder)
    db.commit()
    db.refresh(folder)
    return folder


@router.get("/document-drive-folders/{folder_id}", response_model=schemas.DocumentDriveFolderSchema)
def get_drive_folder(folder_id: int, db: Session = Depends(get_db)):
    return _get_folder_or_404(db, folder_id)


@router.patch("/document-drive-folders/{folder_id}", response_model=schemas.DocumentDriveFolderSchema)
def update_drive_folder(
    folder_id: int,
    payload: schemas.DocumentDriveFolderUpdate,
    db: Session = Depends(get_db)
):
    folder = _get_folder_or_404(db, folder_id)
    if payload.name is not None and payload.name.strip():
        folder.name = payload.name.strip()
    elif payload.name is not None and not payload.name.strip():
        raise HTTPException(status_code=400, detail="Folder name is required")
    if payload.description is not None:
        folder.description = payload.description
    db.commit()
    db.refresh(folder)
    return folder


def _recursive_delete_folder(db: Session, folder):
    """Recursively delete subfolders. Files that are still linked to a contract version are detached
    (``folder_id`` cleared); other files are removed from disk and the ``files`` table."""
    for child in list(folder.children):
        _recursive_delete_folder(db, child)
    for f in list(folder.files):
        if (
            db.query(models.DocumentVersion)
            .filter(models.DocumentVersion.file_id == f.id)
            .count()
            > 0
        ):
            f.folder_id = None
            continue
        _safe_unlink(f.file_path)
        db.delete(f)
    db.delete(folder)


@router.delete("/document-drive-folders/{folder_id}")
def delete_drive_folder(folder_id: int, db: Session = Depends(get_db)):
    folder = _get_folder_or_404(db, folder_id)
    _recursive_delete_folder(db, folder)
    db.commit()
    return {"detail": "Folder and all its contents deleted"}


@router.get("/document-drive-folders/{folder_id}/files", response_model=List[schemas.FileSchema])
def list_drive_folder_files(folder_id: int, db: Session = Depends(get_db)):
    _get_folder_or_404(db, folder_id)
    return (
        db.query(models.File)
        .filter(models.File.folder_id == folder_id)
        .order_by(models.File.created_at.desc())
        .all()
    )


@router.post("/document-drive-folders/{folder_id}/files/chunks", response_model=schemas.DocumentDriveChunkUploadResponse)
async def upload_drive_folder_file_chunk(
    background_tasks: BackgroundTasks,
    folder_id: int,
    upload_id: str = Form(...),
    chunk_index: int = Form(...),
    total_chunks: int = Form(...),
    filename: str = Form(...),
    mime_type: Optional[str] = Form(None),
    chunk: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    folder = _get_folder_or_404(db, folder_id)
    if total_chunks <= 0:
        raise HTTPException(status_code=400, detail="total_chunks must be greater than 0")
    if chunk_index < 0 or chunk_index >= total_chunks:
        raise HTTPException(status_code=400, detail="Invalid chunk_index")

    _cleanup_expired_upload_sessions()
    upload_key = upload_id.strip()
    if not upload_key:
        raise HTTPException(status_code=400, detail="upload_id is required")

    with _DRIVE_UPLOAD_LOCK:
        session = _DRIVE_UPLOAD_SESSIONS.get(upload_key)
        if session is None:
            chunk_dir = os.path.join(DOCUMENT_DRIVE_CHUNKS_DIR, upload_key)
            os.makedirs(chunk_dir, exist_ok=True)
            session = {
                "drive_folder_id": folder_id,
                "total_chunks": total_chunks,
                "filename": filename.strip() or "file",
                "content_type": mime_type,
                "received": set(),
                "chunk_dir": chunk_dir,
                "started_at": time.time(),
            }
            _DRIVE_UPLOAD_SESSIONS[upload_key] = session
        elif session["drive_folder_id"] != folder_id:
            raise HTTPException(status_code=400, detail="upload_id already in use by another folder")
        elif session["total_chunks"] != total_chunks:
            raise HTTPException(status_code=400, detail="Chunk configuration changed during upload")
        elif session["filename"] != filename.strip():
            raise HTTPException(status_code=400, detail="Filename changed during upload")

        chunk_path = os.path.join(session["chunk_dir"], f"{chunk_index:06d}.part")
        if chunk.size:
            chunk.file.seek(0)
        with open(chunk_path, "wb") as f:
            shutil.copyfileobj(chunk.file, f)
        session["received"].add(chunk_index)

        if len(session["received"]) < session["total_chunks"]:
            return schemas.DocumentDriveChunkUploadResponse(
                upload_id=upload_key,
                complete=False,
                file=None,
            )

        for i in range(session["total_chunks"]):
            if i not in session["received"] or not os.path.exists(
                os.path.join(session["chunk_dir"], f"{i:06d}.part")
            ):
                return schemas.DocumentDriveChunkUploadResponse(
                    upload_id=upload_key,
                    complete=False,
                    file=None,
                )

        resolved_filename = _safe_filename(session["filename"])
        final_path = _resolve_drive_file_path(folder_id, resolved_filename, upload_key)
        os.makedirs(os.path.dirname(final_path), exist_ok=True)

        total_bytes = 0
        with open(final_path, "wb") as out:
            for i in range(session["total_chunks"]):
                part_path = os.path.join(session["chunk_dir"], f"{i:06d}.part")
                with open(part_path, "rb") as part_file:
                    while True:
                        data = part_file.read(1024 * 1024)
                        if not data:
                            break
                        out.write(data)
                        total_bytes += len(data)

        _DRIVE_UPLOAD_SESSIONS.pop(upload_key, None)
        shutil.rmtree(session["chunk_dir"], ignore_errors=True)

        ext = os.path.splitext(resolved_filename)[1].lstrip(".").lower()
        record = models.File(
            original_filename=session["filename"],
            file_path=final_path,
            file_type=ext or None,
            content_type=session["content_type"],
            size_bytes=total_bytes,
            folder_id=folder_id,
            upload_id=upload_key,
        )
        db.add(record)
        db.commit()
        db.refresh(record)

        background_tasks.add_task(_run_chunk_file_background, record.id)

        return schemas.DocumentDriveChunkUploadResponse(
            upload_id=upload_key,
            complete=True,
            file=record,
        )


@router.get("/document-drive-files/{file_id}/download")
def download_drive_file(file_id: int, db: Session = Depends(get_db)):
    drive_file = db.query(models.File).filter(models.File.id == file_id).first()
    if not drive_file:
        raise HTTPException(status_code=404, detail="File not found")
    if not os.path.isfile(drive_file.file_path):
        raise HTTPException(status_code=404, detail="Stored file not found")
    return FileResponse(
        path=drive_file.file_path,
        filename=drive_file.original_filename,
        media_type=drive_file.content_type or "application/octet-stream",
    )


@router.get("/document-drive-files/{file_id}/preview")
def preview_drive_file(file_id: int, db: Session = Depends(get_db)):
    drive_file = db.query(models.File).filter(models.File.id == file_id).first()
    if not drive_file:
        raise HTTPException(status_code=404, detail="File not found")
    if not os.path.isfile(drive_file.file_path):
        raise HTTPException(status_code=404, detail="Stored file not found")
    ct = (drive_file.content_type or "").lower()
    if not ct or ct == "application/octet-stream":
        ext = (drive_file.file_type or "").lower()
        if not ext and drive_file.original_filename and "." in drive_file.original_filename:
            ext = drive_file.original_filename.rsplit(".", 1)[-1].lower()
        if ext == "pdf":
            ct = "application/pdf"
        elif ext in ("docx", "doc"):
            ct = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                if ext == "docx"
                else "application/msword"
            )

    if ct.startswith("image/") or ct == "application/pdf" or ct.startswith("text/"):
        return FileResponse(
            path=drive_file.file_path,
            media_type=drive_file.content_type or "application/octet-stream",
        )
    if ct in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        try:
            with open(drive_file.file_path, "rb") as f:
                result = mammoth.convert_to_html(f)
            _css = (
                "html,body{background:#fff;color:#1a1a1a;"
                "font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;"
                "padding:24px;max-width:800px;margin:auto;line-height:1.7;font-size:14px}"
                "h1,h2,h3,h4,h5,h6{color:#111;margin-top:1.4em}"
                "p{margin:0.6em 0}"
                "table{border-collapse:collapse;width:100%}"
                "td,th{border:1px solid #ddd;padding:6px 10px}"
                "a{color:#2563eb}"
            )
            html_out = (
                f"<html><head><meta charset='utf-8'><style>{_css}</style></head>"
                f"<body>{result.value}</body></html>"
            )
            return HTMLResponse(content=html_out)
        except Exception:
            raise HTTPException(status_code=415, detail="Cannot preview this file type")
    raise HTTPException(status_code=415, detail="Preview not supported for this file type")


@router.patch("/document-drive-files/{file_id}/rename", response_model=schemas.FileSchema)
def rename_drive_file(file_id: int, payload: dict, db: Session = Depends(get_db)):
    drive_file = db.query(models.File).filter(models.File.id == file_id).first()
    if not drive_file:
        raise HTTPException(status_code=404, detail="File not found")
    new_name = (payload.get("original_filename") or "").strip()
    if not new_name:
        raise HTTPException(status_code=400, detail="Filename is required")
    drive_file.original_filename = new_name
    if "." in new_name:
        drive_file.file_type = new_name.rsplit(".", 1)[-1].lower()
    db.commit()
    db.refresh(drive_file)
    return drive_file


@router.patch("/document-drive-files/{file_id}/move", response_model=schemas.FileSchema)
def move_drive_file(file_id: int, payload: dict, db: Session = Depends(get_db)):
    drive_file = db.query(models.File).filter(models.File.id == file_id).first()
    if not drive_file:
        raise HTTPException(status_code=404, detail="File not found")
    target_folder_id = payload.get("folder_id")
    if not target_folder_id:
        raise HTTPException(status_code=400, detail="folder_id is required")
    target_folder = db.query(models.DocumentDriveFolder).filter(models.DocumentDriveFolder.id == target_folder_id).first()
    if not target_folder:
        raise HTTPException(status_code=404, detail="Target folder not found")
    if drive_file.folder_id is not None:
        src_folder = db.query(models.DocumentDriveFolder).filter(
            models.DocumentDriveFolder.id == drive_file.folder_id
        ).first()
        if src_folder and src_folder.drive_id != target_folder.drive_id:
            raise HTTPException(
                status_code=400,
                detail="Cannot move a file to a folder in another document drive.",
            )
    if drive_file.folder_id == target_folder_id:
        raise HTTPException(status_code=400, detail="File is already in that folder")
    drive_file.folder_id = target_folder_id
    db.commit()
    db.refresh(drive_file)
    return drive_file


@router.patch("/document-drive-folders/{folder_id}/move", response_model=schemas.DocumentDriveFolderSchema)
def move_drive_folder(folder_id: int, payload: dict, db: Session = Depends(get_db)):
    folder = _get_folder_or_404(db, folder_id)
    target_parent_id = payload.get("parent_id")
    if target_parent_id is not None:
        if target_parent_id == folder_id:
            raise HTTPException(status_code=400, detail="Cannot move folder into itself")
        target = db.query(models.DocumentDriveFolder).filter(
            models.DocumentDriveFolder.id == target_parent_id,
            models.DocumentDriveFolder.drive_id == folder.drive_id,
        ).first()
        if not target:
            raise HTTPException(status_code=404, detail="Target parent folder not found")
        # prevent moving into own descendant
        ancestor = target
        while ancestor:
            if ancestor.id == folder_id:
                raise HTTPException(status_code=400, detail="Cannot move folder into its own descendant")
            ancestor = ancestor.parent
    folder.parent_id = target_parent_id
    db.commit()
    db.refresh(folder)
    return folder


@router.get("/document-drive-files/{file_id}/chunks", response_model=List[schemas.DocumentChunkSchema])
def get_drive_file_chunks(file_id: int, db: Session = Depends(get_db)):
    drive_file = db.query(models.File).filter(models.File.id == file_id).first()
    if not drive_file:
        raise HTTPException(status_code=404, detail="File not found")
    return (
        db.query(models.DocumentChunk)
        .filter(models.DocumentChunk.file_id == file_id)
        .order_by(models.DocumentChunk.chunk_index)
        .all()
    )


@router.get("/versions/{version_id}/chunks", response_model=List[schemas.DocumentChunkSchema])
def get_version_chunks(version_id: int, db: Session = Depends(get_db)):
    version = db.query(models.DocumentVersion).filter(models.DocumentVersion.id == version_id).first()
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")
    if not version.file_id:
        return []
    return (
        db.query(models.DocumentChunk)
        .filter(models.DocumentChunk.file_id == version.file_id)
        .order_by(models.DocumentChunk.chunk_index)
        .all()
    )


@router.delete("/document-drive-files/{file_id}")
def delete_drive_file(file_id: int, db: Session = Depends(get_db)):
    drive_file = db.query(models.File).filter(models.File.id == file_id).first()
    if not drive_file:
        raise HTTPException(status_code=404, detail="File not found")
    if (
        db.query(models.DocumentVersion)
        .filter(models.DocumentVersion.file_id == file_id)
        .count()
        > 0
    ):
        raise HTTPException(
            status_code=409,
            detail="This file is linked to a contract document version; remove or change that version first.",
        )
    _safe_unlink(drive_file.file_path)
    db.delete(drive_file)
    db.commit()
    return {"detail": "File deleted"}
